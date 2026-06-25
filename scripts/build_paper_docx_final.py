from __future__ import annotations

from datetime import date
import hashlib
import re
from pathlib import Path
from shutil import copy2

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
STUDENT_ID = "1030424322"
STUDENT_NAME = "解宝赛"
SUBMIT_DIR = ROOT / f"{STUDENT_ID}-{STUDENT_NAME}"
DOCX_PATH = SUBMIT_DIR / "期末大作业论文.docx"
REPORT_COPY = ROOT / "reports" / "期末大作业论文.docx"

TITLE = "基于传染病动力学模型与机器学习方法的疫情传播预测研究"
COURSE = "计算机建模技术"
TOPIC = "题目一：疫情传播模型与预测模型方法"

REGION_SLUG = "china_mainland"
REGION_NAME = "中国大陆"
RECOVERED_END = pd.Timestamp("2021-08-04")
BOOKMARK_ID = 1


def read_csv(rel_path: str) -> pd.DataFrame:
    path = ROOT / rel_path
    return pd.read_csv(path) if path.exists() else pd.DataFrame()


def set_font(run, size: float = 12, east_asia: str = "宋体", ascii_font: str = "Times New Roman", bold: bool | None = None) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def configure_section(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def configure_styles(doc: Document) -> None:
    configure_section(doc.sections[0])
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    body = styles["Body Text"]
    body.font.name = "Times New Roman"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    body._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    body._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    body.font.size = Pt(12)
    body.paragraph_format.first_line_indent = Pt(24)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    h1 = styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(16)
    h1.paragraph_format.line_spacing = 1.5

    h2 = styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(7)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.5

    h3 = styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h3._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(0)
    h3.paragraph_format.space_after = Pt(0)
    h3.paragraph_format.line_spacing = 1.5


def disable_open_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is not None:
        settings.remove(update)


def set_page_number_start(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.append(pg)
    pg.set(qn("w:start"), str(start))


def add_page_field(paragraph, roman: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* ROMAN" if roman else "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_font(run, 10.5)


def bookmark_name(text: str) -> str:
    return "bm_" + hashlib.md5(text.encode("utf-8")).hexdigest()[:16]


def add_hyperlink_run(paragraph, text: str, anchor: str | None = None, url: str | None = None) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    if anchor:
        hyperlink.set(qn("w:anchor"), anchor)
    if url:
        rid = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink.set(qn("r:id"), rid)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_heading(doc: Document, text: str, level: int) -> None:
    global BOOKMARK_ID
    p = doc.add_paragraph(style=f"Heading {level}")
    p.text = ""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(BOOKMARK_ID))
    start.set(qn("w:name"), bookmark_name(text))
    p._p.append(start)
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 14 if level == 2 else 12, east_asia="黑体", bold=True)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(BOOKMARK_ID))
    p._p.append(end)
    BOOKMARK_ID += 1


def add_p(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_font(run, 12)


def add_center(doc: Document, text: str, size: float = 12, bold: bool = False, east_asia: str = "宋体") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, size, east_asia=east_asia, bold=bold)


def add_no_indent(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, 12, bold=bold)


def add_formula(doc: Document, formula: str, number: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    omath = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = formula
    math_run.append(math_text)
    omath.append(math_run)
    p._p.append(omath)
    run = p.add_run(f"    {number}")
    set_font(run, 12)


def add_reference_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    url_pattern = re.compile(r"https?://[^\s,，]+")
    pos = 0
    for match in url_pattern.finditer(text):
        if match.start() > pos:
            run = p.add_run(text[pos : match.start()])
            set_font(run, 12)
        url = match.group(0).rstrip(".")
        trailing = match.group(0)[len(url) :]
        add_hyperlink_run(p, url, url=url)
        if trailing:
            run = p.add_run(trailing)
            set_font(run, 12)
        pos = match.end()
    if pos < len(text):
        run = p.add_run(text[pos:])
        set_font(run, 12)


def fmt_value(value, col: str) -> str:
    if pd.isna(value):
        return ""
    int_cols = {
        "days",
        "train_rows",
        "test_rows",
        "forecast_day",
        "fit_days",
        "effective_population",
        "negative_daily_values_clipped",
        "missing_days",
        "final_cumulative_confirmed",
        "total_daily_confirmed",
        "peak_daily_confirmed",
    }
    metric_cols = {
        "RMSE",
        "MAE",
        "R2",
        "beta",
        "gamma",
        "R0_beta_over_gamma",
        "mean_beta",
        "mean_gamma",
        "mean_R0_beta_over_gamma",
        "final_R0_beta_over_gamma",
    }
    if col in int_cols:
        return f"{float(value):,.0f}"
    if col in metric_cols:
        return f"{float(value):,.3f}"
    if col == "prediction":
        return f"{float(value):,.2f}"
    if isinstance(value, (float, int)):
        return f"{float(value):,.3f}"
    return str(value)


def shade(cell, fill: str = "F2F2F2") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width))


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER, size: float = 10.5) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.text = ""
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_font(run, size, bold=bold)


def add_table_caption(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(caption)
    set_font(run, 10.5, east_asia="黑体", bold=True)


def add_figure_caption(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(caption)
    set_font(run, 10.5, east_asia="黑体", bold=True)


def add_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str], widths: list[int], max_rows: int | None = None) -> None:
    if df.empty:
        add_no_indent(doc, "暂无可用结果。")
        return
    use = df[[c for c in columns if c in df.columns]].copy()
    if max_rows is not None:
        use = use.head(max_rows)
    table = doc.add_table(rows=1, cols=len(use.columns))
    table.style = "Table Grid"
    set_table_geometry(table, widths[: len(use.columns)])
    for idx, header in enumerate(headers[: len(use.columns)]):
        shade(table.rows[0].cells[idx])
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    left_cols = {"target", "model", "script", "description", "target_column", "evaluation_scope"}
    for _, row in use.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(use.columns):
            align = WD_ALIGN_PARAGRAPH.LEFT if col in left_cols else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[idx], fmt_value(row[col], col), align=align)
    set_table_geometry(table, widths[: len(use.columns)])


def add_figure(doc: Document, rel_path: str, caption: str, width_cm: float = 15.2) -> None:
    path = ROOT / rel_path
    if not path.exists():
        add_no_indent(doc, f"图片缺失：{rel_path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_figure_caption(doc, caption)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    add_center(doc, "期末大作业论文", size=22, bold=True, east_asia="黑体")
    for _ in range(3):
        doc.add_paragraph()
    add_center(doc, TITLE, size=18, bold=True, east_asia="黑体")
    for _ in range(6):
        doc.add_paragraph()
    cover_items = [
        f"课程名称：{COURSE}",
        f"作业题目：{TOPIC}",
        f"学号：{STUDENT_ID}",
        f"姓名：{STUDENT_NAME}",
        f"完成日期：{date.today().isoformat()}",
    ]
    for item in cover_items:
        add_center(doc, item, size=14)


def add_toc(doc: Document) -> None:
    add_heading(doc, "目录", 1)
    items = [
        "摘要",
        "一、引言",
        "二、数据来源与预处理",
        "三、实验设置",
        "四、疫情传播规律与区域趋势分析",
        "五、固定参数 SIR 模型",
        "六、时变参数 SIR 模型",
        "七、机器学习每日新增预测模型",
        "八、同窗口舱室变量对比实验",
        "九、未来 7 天和 14 天预测",
        "十、模型对比分析",
        "十一、结论与不足",
        "参考文献",
    ]
    page_hints = ["I", "1", "2", "3", "4", "8", "10", "12", "16", "18", "20", "21", "22"]
    for item, page_hint in zip(items, page_hints):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        add_hyperlink_run(p, item, anchor=bookmark_name(item))
        dots = "." * max(6, 34 - len(item))
        run = p.add_run(f" {dots} {page_hint}")
        set_font(run, 12)


def date_range_text(df: pd.DataFrame) -> str:
    if df.empty or "date" not in df.columns:
        return "未知"
    dates = pd.to_datetime(df["date"])
    return f"{dates.min().date()} 至 {dates.max().date()}"


def metric_slice(metrics: pd.DataFrame, target: str | None = None, model: str | None = None) -> pd.DataFrame:
    use = metrics.copy()
    if target is not None and "target" in use.columns:
        use = use[use["target"].eq(target)]
    if model is not None and "model" in use.columns:
        use = use[use["model"].eq(model)]
    return use


def best_row(metrics: pd.DataFrame, target: str) -> pd.Series | None:
    use = metric_slice(metrics, target)
    if use.empty:
        return None
    return use.sort_values(["RMSE", "MAE"]).iloc[0]


def add_abstract(doc: Document, processed: pd.DataFrame, ml_metrics: pd.DataFrame) -> None:
    add_heading(doc, "摘要", 1)
    final_confirmed = int(processed["cumulative_confirmed"].dropna().iloc[-1]) if "cumulative_confirmed" in processed else 0
    final_deaths = int(processed["cumulative_deaths"].dropna().iloc[-1]) if "cumulative_deaths" in processed else 0
    best_confirmed = best_row(ml_metrics, "confirmed")
    best_text = ""
    if best_confirmed is not None:
        best_text = f"在每日新增确诊测试集中，{best_confirmed['model']} 的 RMSE 为 {fmt_value(best_confirmed['RMSE'], 'RMSE')}、MAE 为 {fmt_value(best_confirmed['MAE'], 'MAE')}、R² 为 {fmt_value(best_confirmed['R2'], 'R2')}。"
    add_p(
        doc,
        f"本研究围绕“疫情传播模型与预测模型方法”，基于 JHU CSSE COVID-19 时间序列数据，以{REGION_NAME}为分析对象，"
        f"完成数据预处理、区域传播趋势分析、SIR 模型拟合和机器学习短期预测。China 口径中排除 Hong Kong、Macau、Taiwan 和 Unknown；"
        f"样本时间范围为 {date_range_text(processed)}，末期累计确诊 {final_confirmed:,} 例，累计死亡 {final_deaths:,} 例。"
        "模型部分构建固定参数 SIR、时变参数 SIR，并比较 NaivePersistence、LinearRegression、RandomForest 和 GradientBoosting 四类机器学习模型。"
        "其中 TimeVaryingSIR 作为历史描述性重构模型解释参数变化，不作为普通测试集预测结果解读。"
        f"{best_text}"
        "在中国大陆后期每日新增确诊预测中，NaivePersistence 取得最低 RMSE，说明该阶段序列具有较强短期惯性；复杂模型并不一定稳定优于简单基线。"
        "SIR 类模型具有较强机制解释能力；deaths 和 recovered 序列受低计数波动、统计修正和 recovered 停止可靠更新影响较大。"
    )
    add_no_indent(doc, "关键词：疫情传播；SIR 模型；时变参数；机器学习；GradientBoosting；时间序列预测；区域传播趋势")


def build_document() -> None:
    processed = read_csv(f"data/processed/{REGION_SLUG}_timeseries.csv")
    if processed.empty:
        raise FileNotFoundError(f"缺少 data/processed/{REGION_SLUG}_timeseries.csv，请先运行 src/run_all.py。")
    processed["date"] = pd.to_datetime(processed["date"])
    quality = read_csv(f"data/processed/{REGION_SLUG}_data_quality.csv")
    province = read_csv(f"outputs/metrics/{REGION_SLUG}_province_trend_summary.csv")
    sir_params = read_csv(f"outputs/metrics/{REGION_SLUG}_sir_parameters.csv")
    sir_metrics = read_csv(f"outputs/metrics/{REGION_SLUG}_sir_metrics.csv")
    tv_params = read_csv(f"outputs/metrics/{REGION_SLUG}_time_varying_sir_parameters.csv")
    tv_metrics = read_csv(f"outputs/metrics/{REGION_SLUG}_time_varying_sir_metrics.csv")
    ml_metrics = read_csv(f"outputs/metrics/{REGION_SLUG}_ml_metrics.csv")
    comp_metrics = read_csv(f"outputs/metrics/{REGION_SLUG}_compartment_ml_metrics.csv")
    future7 = read_csv(f"outputs/predictions/{REGION_SLUG}_ml_future_7d_confirmed_deaths.csv")
    future14 = read_csv(f"outputs/predictions/{REGION_SLUG}_ml_future_14d_confirmed_deaths.csv")

    doc = Document()
    configure_styles(doc)
    disable_open_update_fields(doc)

    add_cover(doc)

    abstract_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(abstract_section)
    abstract_section.footer.is_linked_to_previous = False
    add_page_field(abstract_section.footer.paragraphs[0], roman=True)
    set_page_number_start(abstract_section, 1)
    add_abstract(doc, processed, ml_metrics)
    doc.add_page_break()
    add_toc(doc)

    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(body_section)
    body_section.footer.is_linked_to_previous = False
    add_page_field(body_section.footer.paragraphs[0], roman=False)
    set_page_number_start(body_section, 1)

    add_heading(doc, "一、引言", 1)
    add_p(
        doc,
        "新冠疫情数据既具有传染病动力学特征，也具有典型时间序列特征。累计确诊、累计死亡和累计治愈能够反映疫情总规模，"
        "每日新增则更直接地反映短期传播强度、集中补报、检测能力变化和防控措施影响。课程大作业要求不仅给出预测结果，还要说明数据来源、模型假设、评价指标和可视化结论。"
    )
    add_p(
        doc,
        "因此，本研究采用两条互补路线：一是使用 SIR 及其时变参数形式解释传播机制；二是使用多种机器学习回归模型学习每日新增序列的短期规律。"
        "前者强调 β、γ、R0 等参数的流行病学含义，后者强调滞后特征、滚动统计特征与日历特征对短期预测的贡献。"
    )
    add_p(
        doc,
        "与只展示单一模型不同，本研究保留 Naive persistence baseline，并加入 LinearRegression、RandomForestRegressor 和 GradientBoostingRegressor。"
        "Naive baseline 表示‘明天等于今天’的简单惯性假设，是判断复杂模型是否真正有效的最低参照；线性回归提供可解释的线性基线；随机森林和梯度提升树用于刻画非线性关系。"
    )

    add_heading(doc, "二、数据来源与预处理", 1)
    add_p(
        doc,
        "数据来源为 Johns Hopkins University Center for Systems Science and Engineering 公开维护的 COVID-19 Data Repository。"
        "原始文件包括 time_series_covid19_confirmed_global.csv、time_series_covid19_deaths_global.csv 和 time_series_covid19_recovered_global.csv。"
        f"本研究保留 JHU 数据源，但将 China 口径进一步限定为{REGION_NAME}：在国家汇总和省级面板中排除 Hong Kong、Macau、Taiwan 和 Unknown。"
        "这样做的目的不是否认这些地区的数据价值，而是为了避免大陆省级趋势图和 Top 10 省份对比中混入不同统计口径。"
    )
    add_p(
        doc,
        "预处理首先将 JHU 宽表日期列转换为标准长时间序列，并按日期升序排列。累计值转换为每日新增值时，使用相邻日期差分；若出现负新增值，"
        "按 0 截断，以处理历史修正或集中回补带来的异常负数。该处理会在数据质量表中记录，而不是在论文中伪造平滑结果。"
    )
    add_formula(doc, "daily_new_t = max(cumulative_t - cumulative_{t-1}, 0)", "（1）")
    add_p(
        doc,
        f"需要特别说明的是，JHU recovered 全球时间序列的可靠结束日期为 {RECOVERED_END.date()}。因此 confirmed 和 deaths 使用完整样本区间 "
        f"{date_range_text(processed)}，而 recovered 只在可靠区间内训练、测试和短期外推；正文未来预测不把 recovered 与 2023 年 confirmed/deaths 结果混在同一张明细表中。"
    )
    add_table_caption(doc, "表 1 数据质量摘要")
    add_table(doc, quality, ["target", "missing_days", "negative_daily_values_clipped"], ["变量", "缺失天数", "负新增截断次数"], [2600, 2200, 2600])

    add_heading(doc, "三、实验设置", 1)
    train_info = ml_metrics[ml_metrics["target"].eq("confirmed")].head(1)
    if not train_info.empty:
        r = train_info.iloc[0]
        split_text = f"每日新增确诊训练期为 {r['train_start_date']} 至 {r['train_end_date']}，测试期为 {r['test_start_date']} 至 {r['test_end_date']}。"
    else:
        split_text = "训练集和测试集均按时间顺序划分，前 80% 为训练集，后 20% 为测试集。"
    add_p(
        doc,
        f"实验数据范围为 {date_range_text(processed)}，时间序列划分不随机打乱。{split_text}"
        "所有机器学习模型使用相同的特征矩阵和相同的训练/测试时间窗口。评价指标包括 RMSE、MAE 和 R²；其中 R² 为负时，表示模型在该测试集上差于使用测试集均值的基线。"
    )
    add_formula(doc, "RMSE = sqrt((1/n) Σ(y_i - ŷ_i)^2)", "（2）")
    add_formula(doc, "MAE = (1/n) Σ|y_i - ŷ_i|", "（3）")
    add_formula(doc, "R² = 1 - Σ(y_i - ŷ_i)^2 / Σ(y_i - ȳ)^2", "（4）")
    params = pd.DataFrame(
        [
            ("NaivePersistence", "ŷ(t)=y(t-1)", "-", "-", "-", "-"),
            ("LinearRegression", "线性最小二乘回归", "-", "-", "-", "-"),
            ("RandomForest", "随机森林回归", "500", "-", "10", "42"),
            ("GradientBoosting", "梯度提升回归树", "500", "0.03", "2", "42"),
        ],
        columns=["model", "description", "n_estimators", "learning_rate", "max_depth", "random_state"],
    )
    add_table_caption(doc, "表 2 机器学习模型与主要参数")
    add_table(doc, params, ["model", "description", "n_estimators", "learning_rate", "max_depth", "random_state"], ["模型", "说明", "树数量", "学习率", "最大深度", "随机种子"], [1600, 2800, 1000, 1000, 1000, 1000])
    add_p(
        doc,
        "程序输出主要保存在 outputs 文件夹：figures 存放趋势图、拟合图、预测图和误差对比图；metrics 存放 SIR 参数、ML 指标、区域 Top 10 表和合并指标表；"
        "predictions 存放测试集预测结果以及未来 7 天、14 天预测 CSV。Word 论文中的图表均引用这些程序输出结果。"
    )

    add_heading(doc, "四、疫情传播规律与区域趋势分析", 1)
    add_p(
        doc,
        "从累计趋势看，确诊和死亡序列总体呈阶梯式增长，说明疫情统计并非稳定连续变化，而是受到传播高峰、检测能力、统计口径和集中补报共同影响。"
        "每日新增曲线比累计曲线更能暴露短期冲击，因此后续机器学习预测以每日新增为主要目标。"
    )
    add_figure(doc, f"outputs/figures/{REGION_SLUG}_cumulative_trends.png", "图 1 中国大陆累计确诊、死亡、治愈趋势")
    add_figure(doc, f"outputs/figures/{REGION_SLUG}_daily_trends.png", "图 2 中国大陆每日新增确诊、死亡、治愈趋势")
    add_p(
        doc,
        "区域传播趋势部分使用省级面板数据绘制累计确诊趋势、每日新增趋势和 Top 10 省份对比。由于 Hong Kong、Macau、Taiwan 和 Unknown 已被排除，"
        "这些图反映的是中国大陆省级记录。Top 10 省份代表累计确诊规模较大的地区，但不等同于风险率；若要比较风险率，还需要按人口规模进行标准化。"
    )
    add_figure(doc, f"outputs/figures/{REGION_SLUG}_province_cumulative_confirmed_trends.png", "图 3 各省累计确诊趋势")
    add_figure(doc, f"outputs/figures/{REGION_SLUG}_province_daily_confirmed_trends.png", "图 4 各省每日新增确诊趋势")
    add_figure(doc, f"outputs/figures/{REGION_SLUG}_province_top10_confirmed_comparison.png", "图 5 累计确诊 Top 10 省份对比")
    top10 = province.sort_values("final_cumulative_confirmed", ascending=False).head(10) if not province.empty else province
    add_table_caption(doc, "表 3 中国大陆累计确诊 Top 10 省份")
    add_table(doc, top10, ["province", "final_cumulative_confirmed", "peak_daily_confirmed", "total_daily_confirmed"], ["省份", "期末累计确诊", "单日新增峰值", "每日新增合计"], [2200, 2300, 2300, 2300])

    add_heading(doc, "五、固定参数 SIR 模型", 1)
    add_p(
        doc,
        "SIR 模型将人群划分为易感者 S、感染者 I 和移除者 R。本文用 confirmed - recovered - deaths 近似感染者，用 recovered + deaths 近似移除者。"
        "固定参数 SIR 假设 β 和 γ 在拟合窗口内保持不变，因此更适合解释平均传播强度，而不是精确追踪每一个统计突变点。"
    )
    add_formula(doc, "dS/dt = -βSI/N", "（5）")
    add_formula(doc, "dI/dt = βSI/N - γI", "（6）")
    add_formula(doc, "dR/dt = γI", "（7）")
    add_formula(doc, "R0 = β/γ", "（8）")
    add_p(
        doc,
        "其中 β 表示传播率，γ 表示移除率，R0 可近似衡量传播强度。降低接触率会降低 β，提高隔离、治疗和恢复效率会提高 γ，二者共同决定疫情是否趋于扩散。"
    )
    add_table_caption(doc, "表 4 固定参数 SIR 拟合参数")
    add_table(doc, sir_params, ["fit_start_date", "fit_end_date", "fit_days", "beta", "gamma", "R0_beta_over_gamma", "effective_population"], ["开始", "结束", "天数", "β", "γ", "R0", "有效人口"], [1100, 1100, 800, 1100, 1100, 1200, 2100], max_rows=1)
    add_p(
        doc,
        "表 4 中的有效人口不是中国大陆真实总人口，而是 SIR 早期拟合窗口中的有效传播规模参数，表示参与该阶段传播和统计观测的近似规模。"
        "由于真实疫情传播受到检测能力、统计口径和区域隔离等因素影响，该参数更适合作为拟合尺度解释，而不能直接理解为实际人口规模。"
    )
    add_figure(doc, f"outputs/figures/{REGION_SLUG}_sir_fit.png", "图 6 固定参数 SIR 模型拟合结果")
    add_table_caption(doc, "表 5 固定参数 SIR 模型评价指标")
    add_table(doc, sir_metrics, ["target", "model", "RMSE", "MAE", "R2"], ["变量", "模型", "RMSE", "MAE", "R²"], [2200, 2200, 1800, 1800, 1200])

    add_heading(doc, "六、时变参数 SIR 模型", 1)
    add_p(
        doc,
        "时变参数 SIR 将 β 和 γ 扩展为 β(t) 和 γ(t)，用于刻画传播率和移除率随时间变化的现象。"
        "这类模型可以更好地重构历史曲线，但必须谨慎解释：本文的 TimeVaryingSIR 是基于历史观测反推参数后的描述性拟合/重构模型，不应与测试集预测模型直接比较。"
    )
    add_formula(doc, "dS/dt = -β(t)SI/N", "（9）")
    add_formula(doc, "dI/dt = β(t)SI/N - γ(t)I", "（10）")
    add_formula(doc, "R_t = β(t)/γ(t)", "（11）")
    add_p(
        doc,
        "因此，若表中出现 R² 接近 1.000，并不表示模型拥有完美未来预测能力，而表示它在使用完整历史数据反推参数后能够高度贴合历史状态。"
        "该结果主要服务于解释传播强度的阶段性变化，而不是作为与 Naive、LinearRegression、RandomForest、GradientBoosting 同口径的测试集预测结果。"
    )
    add_table_caption(doc, "表 6 时变参数 SIR 关键参数")
    add_table(doc, tv_params, ["fit_start_date", "fit_end_date", "fit_days", "mean_beta", "mean_gamma", "mean_R0_beta_over_gamma", "final_R0_beta_over_gamma"], ["开始", "结束", "天数", "平均β", "平均γ", "平均R0", "期末R0"], [1050, 1050, 800, 1150, 1150, 1300, 1300], max_rows=1)
    add_figure(doc, f"outputs/figures/{REGION_SLUG}_time_varying_sir_fit.png", "图 7 时变参数 SIR 历史重构结果")
    tv_show = tv_metrics.copy()
    if "evaluation_scope" not in tv_show.columns:
        tv_show["evaluation_scope"] = "descriptive_reconstruction"
    add_table_caption(doc, "表 7 时变参数 SIR 描述性重构指标")
    add_table(doc, tv_show, ["target", "model", "evaluation_scope", "RMSE", "MAE", "R2"], ["变量", "模型", "评价口径", "RMSE", "MAE", "R²"], [1200, 1700, 2400, 1200, 1200, 900])

    add_heading(doc, "七、机器学习每日新增预测模型", 1)
    add_p(
        doc,
        "机器学习部分以每日新增确诊、每日新增死亡和每日新增治愈为目标变量，构造滞后值、滚动均值、滚动中位数、滚动标准差、指数滑动均值、日期序号、星期和月份等特征。"
        "所有特征只使用预测日期之前的信息，避免未来信息泄露。未来多步预测采用递推方式：先预测第 1 天，再将预测值加入历史序列构造第 2 天特征。"
    )
    add_p(
        doc,
        "本节不再只展示 GradientBoosting，而是同时比较 NaivePersistence、LinearRegression、RandomForest 和 GradientBoosting。"
        "其中 GradientBoosting 参数为 n_estimators=500、learning_rate=0.03、max_depth=2、random_state=42；RandomForest 使用 n_estimators=500、max_depth=10、random_state=42。"
    )
    add_figure(doc, f"outputs/figures/{REGION_SLUG}_confirmed_ml_test_prediction.png", "图 8 每日新增确诊测试集预测对比")
    add_figure(doc, f"outputs/figures/{REGION_SLUG}_deaths_ml_test_prediction.png", "图 9 每日新增死亡测试集预测对比")
    split_info = ml_metrics.drop_duplicates("target").set_index("target")
    if {"confirmed", "deaths"}.issubset(split_info.index):
        confirmed_row = split_info.loc["confirmed"]
        deaths_row = split_info.loc["deaths"]
        recovered_row = split_info.loc["recovered"] if "recovered" in split_info.index else None
        recovered_text = ""
        if recovered_row is not None:
            recovered_text = (
                f"recovered 因可靠更新时间截止于 2021-08-04，使用 {int(recovered_row['train_rows'])} 行训练、"
                f"{int(recovered_row['test_rows'])} 行测试，测试期从 {recovered_row['test_start_date']} 开始。"
            )
        add_p(
            doc,
            f"每日新增确诊和死亡使用相同时间窗口：confirmed 训练 {int(confirmed_row['train_rows'])} 行、测试 {int(confirmed_row['test_rows'])} 行，"
            f"测试期从 {confirmed_row['test_start_date']} 开始；deaths 训练 {int(deaths_row['train_rows'])} 行、测试 {int(deaths_row['test_rows'])} 行，"
            f"测试期从 {deaths_row['test_start_date']} 开始。{recovered_text}"
        )
    add_table_caption(doc, "表 8-1 每日新增确诊预测指标")
    add_table(doc, metric_slice(ml_metrics, "confirmed"), ["model", "RMSE", "MAE", "R2"], ["模型", "RMSE", "MAE", "R²"], [3300, 1800, 1800, 1400])
    add_table_caption(doc, "表 8-2 每日新增死亡预测指标")
    add_table(doc, metric_slice(ml_metrics, "deaths"), ["model", "RMSE", "MAE", "R2"], ["模型", "RMSE", "MAE", "R²"], [3300, 1800, 1800, 1400])
    add_table_caption(doc, "表 8-3 每日新增治愈预测指标")
    add_table(doc, metric_slice(ml_metrics, "recovered"), ["model", "RMSE", "MAE", "R2"], ["模型", "RMSE", "MAE", "R²"], [3300, 1800, 1800, 1400])
    add_figure(doc, f"outputs/figures/{REGION_SLUG}_ml_metrics_comparison.png", "图 10 机器学习模型 RMSE、MAE、R² 对比（R² 小于 -5 的柱在图中截断显示，完整数值见表 8-1 至表 8-3）", width_cm=14.8)
    add_p(
        doc,
        "从指标看，confirmed 往往比 deaths 和 recovered 更适合做短期预测，因为确诊序列规模较大、可学习模式更多。deaths 单日数值较低时，少量波动就会造成较大的相对误差；"
        "recovered 还受到 2021-08-04 后停止可靠更新的影响，因此其 R² 可能为负。负 R² 并非程序错误，而是提示模型在该任务上不如简单均值基线。"
    )

    add_heading(doc, "八、同窗口舱室变量对比实验", 1)
    add_p(
        doc,
        "为避免传统模型和机器学习模型比较对象不一致，本文补充同窗口舱室变量实验：将机器学习模型应用于 SIR 相同窗口内的 infected、removed 和 confirmed 三个变量。"
        "该实验仍按时间顺序划分训练集和测试集，用于观察机器学习模型在舱室变量上的外推能力。"
    )
    add_p(
        doc,
        "结果中若 GradientBoosting 的 R² 为负，说明树模型在短窗口后段没有超过均值基线。这通常与样本窗口较短、舱室变量存在明显单调趋势、测试期分布偏离训练期有关。"
        "SIR 模型虽然假设较强，但舱室结构本身提供了机制约束；机器学习模型若缺少政策、人口流动、检测量等外生变量，容易在这类外推任务中表现不稳定。"
    )
    add_figure(doc, f"outputs/figures/{REGION_SLUG}_compartment_ml_test_prediction.png", "图 11 舱室变量同窗口机器学习测试集预测")
    add_table_caption(doc, "表 9 同窗口舱室变量机器学习指标")
    add_table(doc, comp_metrics, ["target", "model", "train_rows", "test_rows", "RMSE", "MAE", "R2"], ["目标", "模型", "训练行", "测试行", "RMSE", "MAE", "R²"], [1000, 1700, 800, 800, 1400, 1400, 900])

    add_heading(doc, "九、未来 7 天和 14 天预测", 1)
    add_p(
        doc,
        "未来预测以 confirmed 为正文重点，deaths 简要展示，recovered 不放入正文长表。原因是 recovered 的可靠更新时间截止于 2021-08-04，若继续外推，"
        "只能解释为可靠区间结束后的短期外推，不代表 2023 年真实治愈趋势。"
    )
    add_p(
        doc,
        "需要说明的是，GradientBoosting 不是 confirmed 测试集上的最优模型，表 8-1 显示 NaivePersistence 的 RMSE 最低。"
        "正文展示 GradientBoosting 的未来预测，是为了观察非平凡模型在滞后特征和滚动特征下给出的趋势变化；最终最优误差模型仍以表 8-1 至表 8-3 为准。"
    )
    gb7 = future7[future7["model"].eq("GradientBoosting")].copy() if not future7.empty else future7
    gb14 = future14[future14["model"].eq("GradientBoosting")].copy() if not future14.empty else future14
    confirmed14 = gb14[gb14["target"].eq("confirmed")]
    add_figure(doc, f"outputs/figures/{REGION_SLUG}_confirmed_ml_future_7d.png", "图 12 每日新增确诊未来 7 天预测")
    add_figure(doc, f"outputs/figures/{REGION_SLUG}_confirmed_ml_future_14d.png", "图 13 每日新增确诊未来 14 天预测")
    add_p(doc, "每日新增死亡的 7 天预测值均接近 0，完整结果保存在 outputs/predictions 中。")
    add_table_caption(doc, "表 10 GradientBoosting 每日新增确诊未来 14 天预测")
    add_table(doc, confirmed14, ["date", "forecast_day", "prediction"], ["日期", "预测天数", "预测值"], [2400, 1800, 2400], max_rows=14)

    add_heading(doc, "十、模型对比分析", 1)
    add_p(
        doc,
        "SIR 模型和机器学习模型回答的问题不同。固定参数 SIR 和时变参数 SIR 适合解释感染者、移除者和累计确诊之间的结构关系，"
        "能够把接触率、移除率和传播强度联系起来；机器学习模型更适合在近期历史模式延续的前提下预测每日新增。"
    )
    add_p(
        doc,
        "从模型对比角度看，NaivePersistence 是最低基线。如果 LinearRegression、RandomForest 或 GradientBoosting 没有明显优于 Naive，"
        "说明复杂模型虽然参数更多，但未必能在当前测试窗口带来稳定收益。GradientBoosting 在 confirmed 上通常具备较强短期拟合能力，"
        "但在 deaths 和 recovered 上可能因为低计数波动、统计稀疏、停止更新或分布漂移而得到负 R²。"
    )
    add_p(
        doc,
        "TimeVaryingSIR 的高 R² 不应与测试集预测模型并列解读。它使用历史信息反推 β(t) 和 γ(t)，本质上更接近历史重构；"
        "而机器学习指标来自时间顺序测试集，目标是检验未参与训练时期的预测效果。区分描述性拟合与预测性评估，是本次结果解释中最重要的修正。"
    )

    add_heading(doc, "十一、结论与不足", 1)
    add_p(
        doc,
        "本研究完成了一个可运行、可复现的疫情传播建模流程：读取 JHU 真实数据，按中国大陆口径进行预处理，生成每日新增和区域趋势图，"
        "拟合固定参数 SIR 与时变参数 SIR，训练多个机器学习模型，并输出测试集指标、未来 7 天和 14 天预测、CSV 结果和 Word 论文。"
    )
    add_p(
        doc,
        "不足主要包括四点：第一，JHU 数据存在历史修正、集中补报和 recovered 后期停止可靠更新问题；第二，模型未引入政策强度、检测量、人口流动和疫苗接种等外生变量；"
        "第三，未来预测采用递推策略，预测期越长误差累积风险越高；第四，省份 Top 10 只比较原始规模，没有进行人口标准化，不能直接解释为风险率。"
    )
    add_p(
        doc,
        "后续若继续改进，可在保持传统模型可解释性的基础上加入 SEIR 或 SEIQR 扩展模型，并引入更多外生变量；机器学习部分可继续比较 XGBoost、LSTM 或 GRU，"
        "但仍应保留 Naive 与线性模型作为判断复杂模型是否真正有效的参照。"
    )

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Johns Hopkins University Center for Systems Science and Engineering. COVID-19 Data Repository by CSSE at Johns Hopkins University[DB/OL]. GitHub, 2020-2023. https://github.com/CSSEGISandData/COVID-19, 访问日期: 2026-06-25.",
        "[2] Dong E, Du H, Gardner L. An interactive web-based dashboard to track COVID-19 in real time[J]. The Lancet Infectious Diseases, 2020, 20(5): 533-534. https://doi.org/10.1016/S1473-3099(20)30120-1.",
        "[3] Kermack W O, McKendrick A G. A contribution to the mathematical theory of epidemics[J]. Proceedings of the Royal Society A, 1927, 115(772): 700-721. https://doi.org/10.1098/rspa.1927.0118.",
        "[4] Hethcote H W. The mathematics of infectious diseases[J]. SIAM Review, 2000, 42(4): 599-653. https://doi.org/10.1137/S0036144500371907.",
        "[5] Anderson R M, May R M. Infectious Diseases of Humans: Dynamics and Control[M]. Oxford: Oxford University Press, 1991. ISBN: 9780198540403.",
        "[6] World Health Organization. Coronavirus disease (COVID-19) pandemic[EB/OL]. https://www.who.int/emergencies/diseases/novel-coronavirus-2019, 访问日期: 2026-06-25.",
        "[7] Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: Machine learning in Python[J]. Journal of Machine Learning Research, 2011, 12: 2825-2830. https://jmlr.org/papers/v12/pedregosa11a.html.",
        "[8] Friedman J H. Greedy function approximation: A gradient boosting machine[J]. Annals of Statistics, 2001, 29(5): 1189-1232. https://doi.org/10.1214/aos/1013203451.",
        "[9] Hastie T, Tibshirani R, Friedman J. The Elements of Statistical Learning[M]. New York: Springer, 2009. ISBN: 9780387848570. https://hastie.su.domains/ElemStatLearn/.",
        "[10] Hyndman R J, Athanasopoulos G. Forecasting: Principles and Practice[M/OL]. 3rd ed. OTexts, 2021. https://otexts.com/fpp3/, 访问日期: 2026-06-25.",
        "[11] James G, Witten D, Hastie T, Tibshirani R. An Introduction to Statistical Learning[M]. 2nd ed. New York: Springer, 2021. ISBN: 9781071614174. https://www.statlearning.com/.",
        "[12] 中华人民共和国国家卫生健康委员会. 疫情通报[EB/OL]. https://www.nhc.gov.cn/xcs/yqtb/list_gzbd.shtml, 访问日期: 2026-06-25.",
    ]
    for ref in refs:
        add_reference_paragraph(doc, ref)

    SUBMIT_DIR.mkdir(exist_ok=True)
    DOCX_PATH.parent.mkdir(exist_ok=True)
    doc.save(DOCX_PATH)
    REPORT_COPY.parent.mkdir(exist_ok=True)
    copy2(DOCX_PATH, REPORT_COPY)
    print(f"Saved DOCX: {DOCX_PATH}")
    print(f"Saved report copy: {REPORT_COPY}")


if __name__ == "__main__":
    build_document()
