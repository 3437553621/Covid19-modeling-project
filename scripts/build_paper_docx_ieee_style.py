from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEST = ROOT / "1030424322-解宝赛" / "期末大作业论文.docx"

TITLE = "基于传染病动力学模型与机器学习方法的疫情传播预测研究"
AUTHOR = "解宝赛"
STUDENT_ID = "1030424322"
COURSE = "计算机建模技术"

INK = RGBColor(0, 0, 0)
MUTED = RGBColor(80, 80, 80)
LIGHT_GRAY = "F2F2F2"
BORDER = "7F7F7F"


def read_csv(path: str) -> pd.DataFrame:
    file_path = ROOT / path
    return pd.read_csv(file_path) if file_path.exists() else pd.DataFrame()


def fmt(value, digits: int = 3) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, (int, float)):
        value = float(value)
        if abs(value) >= 1000:
            return f"{value:,.{digits}f}"
        return f"{value:.{digits}f}"
    return str(value)


def set_run_font(run, size: float = 10.0, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.color.rgb = INK
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_page(section) -> None:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.66)
    section.right_margin = Inches(0.66)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def set_columns(section, num: int = 1, space_dxa: int = 360) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols = cols[0] if cols else OxmlElement("w:cols")
    if cols.getparent() is None:
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_dxa))


def configure_styles(doc: Document) -> None:
    configure_page(doc.sections[0])
    set_columns(doc.sections[0], 1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.03

    body = styles["Body Text"]
    body.font.name = "Times New Roman"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    body.font.size = Pt(10)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Inches(0.18)
    body.paragraph_format.space_after = Pt(2)
    body.paragraph_format.line_spacing = 1.03

    for name, size in [("Heading 1", 10.5), ("Heading 2", 10), ("Heading 3", 10)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = INK
        style.paragraph_format.space_before = Pt(7)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if name == "Heading 1" else WD_ALIGN_PARAGRAPH.LEFT


def add_para(doc: Document, text: str, style: str = "Body Text", indent: bool = True) -> None:
    p = doc.add_paragraph(text, style=style)
    if not indent:
        p.paragraph_format.first_line_indent = Inches(0)


def add_center_para(doc: Document, text: str, size: float, bold: bool = False, italic: bool = False, after: float = 2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)


def add_formula(doc: Document, formula: str, number: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{formula}        {number}")
    run.font.name = "Cambria Math"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    run.font.size = Pt(10)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.0, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.text = ""
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in [("top", 45), ("bottom", 45), ("start", 70), ("end", 70)]:
                elem = tc_mar.find(qn(f"w:{side}"))
                if elem is None:
                    elem = OxmlElement(f"w:{side}")
                    tc_mar.append(elem)
                elem.set(qn("w:w"), str(value))
                elem.set(qn("w:type"), "dxa")


def add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=8.5, bold=True)


def add_figure_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=8.5, bold=True)


def add_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str], widths: list[int], max_rows: int | None = None) -> None:
    if df.empty:
        add_para(doc, "暂无可用结果。", indent=False)
        return
    use = df[[c for c in columns if c in df.columns]].copy()
    if max_rows is not None:
        use = use.head(max_rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], LIGHT_GRAY)
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for _, row in use.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(use.columns):
            align = WD_ALIGN_PARAGRAPH.LEFT if col in {"province", "target", "model", "script", "description"} else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[idx], fmt(row[col]), align=align)
    set_table_geometry(table, widths)


def add_figure(doc: Document, rel: str, caption: str, width: float) -> None:
    path = ROOT / rel
    if not path.exists():
        add_para(doc, f"图片缺失：{rel}", indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_figure_caption(doc, caption)


def full_width_section(doc: Document):
    sec = doc.add_section(WD_SECTION_START.CONTINUOUS)
    configure_page(sec)
    set_columns(sec, 1)
    return sec


def two_col_section(doc: Document):
    sec = doc.add_section(WD_SECTION_START.CONTINUOUS)
    configure_page(sec)
    set_columns(sec, 2, 360)
    return sec


def add_title_block(doc: Document, processed: pd.DataFrame) -> None:
    add_center_para(doc, TITLE, size=17.5, bold=True, after=6)
    add_center_para(doc, f"{AUTHOR}    学号：{STUDENT_ID}", size=10.5, after=1)
    add_center_para(doc, f"{COURSE} 期末大作业 · 题目一：疫情传播模型与预测模型方法", size=9.5, after=5)
    if not processed.empty:
        processed["date"] = pd.to_datetime(processed["date"])
        start = processed["date"].min().date().isoformat()
        end = processed["date"].max().date().isoformat()
        add_center_para(doc, f"数据区间：{start} 至 {end}；数据来源：JHU CSSE COVID-19 Data", size=8.8, italic=True, after=7)


def add_abstract(doc: Document, processed: pd.DataFrame) -> None:
    final_confirmed = int(processed["cumulative_confirmed"].dropna().iloc[-1])
    final_deaths = int(processed["cumulative_deaths"].dropna().iloc[-1])
    final_recovered = int(processed["cumulative_recovered"].dropna().iloc[-1])

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    label = p.add_run("摘要—")
    set_run_font(label, size=9, bold=True)
    body = p.add_run(
        f"本文基于 Johns Hopkins University CSSE COVID-19 全球时间序列数据，对中国疫情传播趋势进行建模与预测研究。"
        f"数据覆盖累计确诊、死亡和治愈序列，期末累计确诊 {final_confirmed:,} 例、累计死亡 {final_deaths:,} 例、可靠治愈累计值为 {final_recovered:,} 例。"
        "研究首先完成累计值到每日新增值的转换，处理负新增与 recovered 后期停止更新问题；随后绘制全国和省份层面的区域传播趋势图；"
        "模型部分构建固定参数 SIR、时变参数 SIR 以及 GradientBoosting 回归模型，完成测试集评价和未来 7/14 天预测。"
        "实验表明，SIR 类模型在舱室变量解释上具有优势，GradientBoosting 更适合短期每日新增预测，但在样本窗口较短或统计口径变化时存在外推风险。"
    )
    set_run_font(body, size=9)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_after = Pt(8)
    label = p.add_run("关键词—")
    set_run_font(label, size=9, bold=True, italic=True)
    body = p.add_run("疫情传播；SIR 模型；时变参数；GradientBoosting；时间序列预测；区域传播趋势；模型评价")
    set_run_font(body, size=9, italic=True)


def section_heading(doc: Document, number: str, title: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.text = ""
    r = p.add_run(f"{number}. {title}")
    set_run_font(r, size=10.5, bold=True)


def build_doc() -> None:
    doc = Document()
    configure_styles(doc)

    processed = read_csv("data/processed/china_all_timeseries.csv")
    quality = read_csv("data/processed/china_all_data_quality.csv")
    province = read_csv("outputs/metrics/china_province_trend_summary.csv")
    sir_params = read_csv("outputs/metrics/china_all_sir_parameters.csv")
    sir_metrics = read_csv("outputs/metrics/china_all_sir_metrics.csv")
    tv_params = read_csv("outputs/metrics/china_all_time_varying_sir_parameters.csv")
    tv_metrics = read_csv("outputs/metrics/china_all_time_varying_sir_metrics.csv")
    ml_metrics = read_csv("outputs/metrics/china_all_ml_metrics.csv")
    comp_metrics = read_csv("outputs/metrics/china_all_compartment_ml_metrics.csv")
    future7 = read_csv("outputs/predictions/china_all_ml_future_7d.csv")
    future14 = read_csv("outputs/predictions/china_all_ml_future_14d.csv")

    add_title_block(doc, processed)
    add_abstract(doc, processed)

    two_col_section(doc)
    section_heading(doc, "I", "引言")
    add_para(doc, "疫情传播数据同时具有传染病动力学特征和时间序列统计特征。累计确诊、死亡和治愈能够反映总体规模，每日新增序列则更能体现短期传播强度、上报波动和政策干预后的阶段性变化。课程大作业不仅要求预测结果，还要求模型假设、评价指标、可视化分析与论文式说明相互一致。")
    add_para(doc, "本文采用“机制模型 + 机器学习模型”的组合思路：固定参数 SIR 用于解释传播率和移除率的基本含义，时变参数 SIR 用于描述传播强度随时间变化，GradientBoosting 则基于滞后特征和滑动统计量完成每日新增预测。同时，为补足题目明确要求的区域传播趋势，本文增加省份层面的累计趋势、每日新增趋势和 Top 10 省份对比。")

    section_heading(doc, "II", "数据来源与预处理")
    add_para(doc, "本文使用 JHU CSSE COVID-19 global time series 数据，核心文件包括累计确诊、累计死亡和累计治愈三类宽表。程序先按 Country/Region 筛选 China，若 Province/State 为空则对全国所有省份行求和；同时保留省份面板用于区域传播趋势分析。")
    add_formula(doc, r"daily_new(t)=max(cumulative(t)-cumulative(t-1),0)", "（1）")
    add_para(doc, "式（1）用于将累计序列转换为每日新增序列。由于真实疫情数据可能出现统计回补或口径修正，累计差分后会产生负新增值，本文将其截断为 0，并在数据质量表中记录负新增修正次数。JHU recovered 序列在后期存在停止更新或终端归零现象，程序将相关日期后的治愈累计值视为缺失，不伪造治愈数据。")

    full_width_section(doc)
    add_table_caption(doc, "表 1  数据质量与可靠区间摘要")
    add_table(doc, quality, ["target", "start_date", "end_date", "reliable_end_date", "days", "negative_daily_values_clipped", "missing_days"], ["指标", "开始日期", "结束日期", "可靠结束", "天数", "负新增截断", "缺失天数"], [900, 1200, 1200, 1200, 800, 1300, 1000])
    two_col_section(doc)

    section_heading(doc, "III", "区域传播趋势分析")
    add_para(doc, "全国趋势图用于观察疫情总体规模和阶段性变化，省份趋势图用于比较不同地区之间的传播差异。本文将 JHU 数据中的 Unknown 视为未分配地区，在汇总 CSV 中保留以保证数据透明，但不将其作为 Top 10 省份排序对象。")

    full_width_section(doc)
    add_figure(doc, "outputs/figures/china_all_cumulative_trends.png", "图 1  全国累计确诊、死亡、治愈趋势", 6.6)
    add_figure(doc, "outputs/figures/china_all_daily_trends.png", "图 2  全国每日新增确诊、死亡、治愈 7 日均值趋势", 6.6)
    add_figure(doc, "outputs/figures/china_province_cumulative_confirmed_trends.png", "图 3  各省累计确诊趋势", 6.6)
    add_figure(doc, "outputs/figures/china_province_daily_confirmed_trends.png", "图 4  各省每日新增确诊 7 日均值趋势", 6.6)
    add_figure(doc, "outputs/figures/china_province_top10_confirmed_comparison.png", "图 5  Top 10 省份累计确诊对比", 5.9)
    named = province[province["province"].astype(str).str.upper() != "UNKNOWN"] if not province.empty else province
    add_table_caption(doc, "表 2  Top 10 省份累计确诊与新增高峰")
    add_table(doc, named, ["province", "final_cumulative_confirmed", "total_daily_confirmed", "peak_daily_confirmed"], ["省份", "最终累计确诊", "每日新增合计", "单日新增峰值"], [1800, 2200, 2200, 2100], max_rows=10)
    two_col_section(doc)

    section_heading(doc, "IV", "固定参数 SIR 模型")
    add_para(doc, "SIR 模型将人群划分为易感者 S、感染者 I 和移除者 R。本文用累计治愈与累计死亡之和近似移除者 R，用累计确诊减去移除者近似当前感染者 I。固定参数 SIR 假设传播率 β 和移除率 γ 在拟合窗口内保持不变。")
    add_formula(doc, r"dS/dt=-βSI/N", "（2）")
    add_formula(doc, r"dI/dt=βSI/N-γI", "（3）")
    add_formula(doc, r"dR/dt=γI", "（4）")
    add_formula(doc, r"R0=β/γ", "（5）")
    add_para(doc, "其中 N 表示有效人口规模，β 反映接触传播强度，γ 反映感染者转入移除状态的速率。降低接触率会降低 β，提高隔离和治疗效率会提高 γ，因此二者共同影响疫情传播趋势。")

    full_width_section(doc)
    add_table_caption(doc, "表 3  固定参数 SIR 拟合参数")
    add_table(doc, sir_params, ["fit_start_date", "fit_end_date", "fit_days", "beta", "gamma", "R0_beta_over_gamma", "effective_population"], ["开始", "结束", "天数", "β", "γ", "R0", "有效人口"], [1050, 1050, 700, 950, 950, 1100, 1900], max_rows=1)
    add_figure(doc, "outputs/figures/china_all_sir_fit.png", "图 6  固定参数 SIR 模型拟合结果", 6.3)
    add_table_caption(doc, "表 4  固定参数 SIR 模型评价指标")
    add_table(doc, sir_metrics, ["target", "model", "RMSE", "MAE", "R2"], ["目标变量", "模型", "RMSE", "MAE", "R2"], [2000, 1500, 1700, 1700, 1200])
    two_col_section(doc)

    section_heading(doc, "V", "时变参数 SIR 模型")
    add_para(doc, "固定参数 SIR 难以描述防控政策、检测能力、医疗资源和公众行为变化带来的阶段差异。时变参数 SIR 将 β 和 γ 扩展为 β(t) 与 γ(t)，从而刻画传播强度和移除效率随时间变化的过程。")
    add_formula(doc, r"dS/dt=-β(t)SI/N", "（6）")
    add_formula(doc, r"dI/dt=β(t)SI/N-γ(t)I", "（7）")
    add_formula(doc, r"Rt=β(t)/γ(t)", "（8）")

    full_width_section(doc)
    add_table_caption(doc, "表 5  时变参数 SIR 关键参数")
    add_table(doc, tv_params, ["fit_start_date", "fit_end_date", "fit_days", "mean_beta", "mean_gamma", "mean_R0_beta_over_gamma", "final_R0_beta_over_gamma"], ["开始", "结束", "天数", "平均β", "平均γ", "平均Rt", "期末Rt"], [1050, 1050, 700, 1050, 1050, 1300, 1300], max_rows=1)
    add_figure(doc, "outputs/figures/china_all_time_varying_sir_fit.png", "图 7  时变参数 SIR 模型拟合结果", 6.3)
    add_table_caption(doc, "表 6  时变参数 SIR 模型评价指标")
    add_table(doc, tv_metrics, ["target", "model", "RMSE", "MAE", "R2"], ["目标变量", "模型", "RMSE", "MAE", "R2"], [1900, 1900, 1700, 1700, 1200])
    two_col_section(doc)

    section_heading(doc, "VI", "GradientBoosting 每日新增预测")
    add_para(doc, "机器学习模型使用 GradientBoostingRegressor，分别对每日新增确诊、每日新增死亡和每日新增治愈建模。特征包括 1 至 14 天及 21、28 天滞后值，3、7、14、28 日滑动均值、滑动中位数、滑动标准差、滑动最大值、指数滑动均值、日期序号、星期和月份等。所有特征只使用预测日期以前的信息，避免未来信息泄露。")
    add_formula(doc, r"ŷ(t)=F(x(t);θ),  x(t)={lag, rolling, ewm, calendar}", "（9）")

    full_width_section(doc)
    add_figure(doc, "outputs/figures/china_all_confirmed_ml_test_prediction.png", "图 8  每日新增确诊测试集预测对比", 6.2)
    add_figure(doc, "outputs/figures/china_all_deaths_ml_test_prediction.png", "图 9  每日新增死亡测试集预测对比", 6.2)
    add_figure(doc, "outputs/figures/china_all_recovered_ml_test_prediction.png", "图 10  每日新增治愈测试集预测对比", 6.2)
    add_table_caption(doc, "表 7  GradientBoosting 每日新增预测指标")
    add_table(doc, ml_metrics, ["target", "model", "train_rows", "test_rows", "RMSE", "MAE", "R2"], ["目标", "模型", "训练行", "测试行", "RMSE", "MAE", "R2"], [1000, 1700, 850, 850, 1450, 1450, 1160])
    add_figure(doc, "outputs/figures/china_all_ml_metrics_comparison.png", "图 11  GradientBoosting 每日新增预测指标对比", 5.7)
    two_col_section(doc)

    section_heading(doc, "VII", "同窗口舱室变量公平对比")
    add_para(doc, "为使传统模型和机器学习模型在同一类变量上可比，本文将 GradientBoosting 应用于 SIR 相同观测窗口内的 infected、removed 和 confirmed 三个目标。该实验仍按时间顺序划分训练集与测试集。结果中的负 R2 被保留，说明该短窗口后段上机器学习模型表现差于均值基线，这反映出样本窗口较短和趋势变化剧烈时树模型外推能力有限。")

    full_width_section(doc)
    add_figure(doc, "outputs/figures/china_all_compartment_ml_test_prediction.png", "图 12  GradientBoosting 舱室变量同窗口测试预测", 6.2)
    add_table_caption(doc, "表 8  GradientBoosting 舱室变量同窗口实验指标")
    add_table(doc, comp_metrics, ["target", "model", "train_rows", "test_rows", "RMSE", "MAE", "R2"], ["目标", "模型", "训练行", "测试行", "RMSE", "MAE", "R2"], [1000, 1700, 850, 850, 1450, 1450, 1160])
    two_col_section(doc)

    section_heading(doc, "VIII", "未来预测与评价指标")
    add_para(doc, "未来预测采用递推方式：每预测一天，就将该预测值加入历史序列，再构造下一天的滞后和滑动特征。本文同时输出未来 7 天和 14 天结果。模型评价采用 RMSE、MAE 和 R2；当 R2 为负时，表示模型在测试集上不如均值基线，本文如实保留。")
    add_formula(doc, r"RMSE=sqrt((1/n)Σ(yi-ŷi)^2)", "（10）")
    add_formula(doc, r"MAE=(1/n)Σ|yi-ŷi|", "（11）")
    add_formula(doc, r"R²=1-Σ(yi-ŷi)^2/Σ(yi-ȳ)^2", "（12）")

    full_width_section(doc)
    add_figure(doc, "outputs/figures/china_all_confirmed_ml_future_7d.png", "图 13  每日新增确诊未来 7 天预测", 6.1)
    add_figure(doc, "outputs/figures/china_all_confirmed_ml_future_14d.png", "图 14  每日新增确诊未来 14 天预测", 6.1)
    add_table_caption(doc, "表 9  未来 7 天预测结果")
    add_table(doc, future7, ["target", "model", "date", "forecast_day", "prediction"], ["目标", "模型", "日期", "预测天数", "预测值"], [1100, 1800, 1550, 1100, 1900], max_rows=21)
    add_table_caption(doc, "表 10  未来 14 天预测结果")
    add_table(doc, future14, ["target", "model", "date", "forecast_day", "prediction"], ["目标", "模型", "日期", "预测天数", "预测值"], [1100, 1800, 1550, 1100, 1900], max_rows=42)
    two_col_section(doc)

    section_heading(doc, "IX", "结论与不足")
    add_para(doc, "本文完成了从真实疫情数据读取、预处理、区域趋势分析、SIR/时变 SIR 建模、GradientBoosting 每日新增预测、未来 7/14 天预测到论文式结果整理的完整流程。固定参数 SIR 结构清楚、参数含义明确；时变参数 SIR 能更好地描述阶段性变化；GradientBoosting 能利用滞后特征完成短期预测，但在舱室变量短窗口实验中表现不稳定。")
    add_para(doc, "不足之处包括：JHU 数据存在回补、修正和 recovered 后期停止更新；模型未显式加入政策干预、人口流动和检测能力等外生变量；递推预测存在误差累积；省份 Top 10 对比只反映原始确诊规模，不能直接等同于人口标准化风险率。后续可加入外生变量、人口标准化指标或更稳健的时序模型进行改进。")

    section_heading(doc, "参考文献", "")
    refs = [
        "[1] Johns Hopkins University Center for Systems Science and Engineering. COVID-19 Data Repository.",
        "[2] W. O. Kermack and A. G. McKendrick, “A contribution to the mathematical theory of epidemics.”",
        "[3] scikit-learn developers. Gradient Boosting regression documentation.",
        "[4] Overleaf. IEEE Conference Template and Journal Article Template Gallery.",
    ]
    for ref in refs:
        add_para(doc, ref, indent=False)

    full_width_section(doc)
    add_table_caption(doc, "表 11  核心脚本与功能对应关系")
    scripts = pd.DataFrame(
        [
            ("src/prepare_data.py", "读取 JHU 数据，生成全国聚合序列和省份面板。"),
            ("src/run_regional_trends.py", "生成全国趋势图、省份趋势图和 Top 10 省份对比图。"),
            ("src/run_sir.py", "拟合固定参数 SIR，输出参数、指标和拟合图。"),
            ("src/run_time_varying_sir.py", "拟合时变参数 SIR，输出每日参数和拟合图。"),
            ("src/run_ml.py", "训练 GradientBoosting 每日新增预测模型并输出 7/14 天预测。"),
            ("src/run_compartment_ml.py", "在 SIR 同窗口上对 infected、removed、confirmed 做机器学习对比。"),
        ],
        columns=["script", "description"],
    )
    add_table(doc, scripts, ["script", "description"], ["脚本", "功能"], [2600, 6000])

    DEST.parent.mkdir(exist_ok=True)
    doc.save(DEST)


if __name__ == "__main__":
    build_doc()
