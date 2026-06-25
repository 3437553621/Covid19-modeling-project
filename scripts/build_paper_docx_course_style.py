from __future__ import annotations

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEST = ROOT / "1030424322-解宝赛" / "期末大作业论文.docx"

TITLE = "基于传染病动力学模型与机器学习方法的疫情传播预测研究"
COURSE = "计算机建模技术"
TOPIC = "题目一：疫情传播模型与预测模型方法"
STUDENT_ID = "1030424322"
STUDENT_NAME = "解宝赛"

CONTENT_WIDTH_CM = 16.0


def read_csv(path: str) -> pd.DataFrame:
    file_path = ROOT / path
    return pd.read_csv(file_path) if file_path.exists() else pd.DataFrame()


def fmt(value, digits: int = 3) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, (float, int)):
        value = float(value)
        if abs(value) >= 1000:
            return f"{value:,.{digits}f}"
        return f"{value:.{digits}f}"
    return str(value)


def fmt_cell(value, col: str) -> str:
    if pd.isna(value):
        return ""
    integer_cols = {
        "days",
        "train_rows",
        "test_rows",
        "forecast_day",
        "negative_daily_values_clipped",
        "missing_days",
        "final_cumulative_confirmed",
        "total_daily_confirmed",
        "peak_daily_confirmed",
        "effective_population",
        "fit_days",
    }
    metric_cols = {"RMSE", "MAE", "R2", "beta", "gamma", "R0_beta_over_gamma", "mean_beta", "mean_gamma", "mean_R0_beta_over_gamma", "final_R0_beta_over_gamma"}
    if col in integer_cols:
        return f"{float(value):,.0f}"
    if col == "prediction":
        return f"{float(value):,.2f}"
    if col in metric_cols:
        return f"{float(value):,.3f}"
    return fmt(value)


def set_run_font(run, size_pt: float = 12, east_asia: str = "宋体", ascii_font: str = "Times New Roman", bold: bool | None = None) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def configure_section(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def configure_styles(doc: Document) -> None:
    configure_section(doc.sections[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    body = doc.styles["Body Text"]
    body.font.name = "Times New Roman"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    body._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    body._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    body.font.size = Pt(12)
    body.paragraph_format.first_line_indent = Pt(24)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(16)
    h1.paragraph_format.line_spacing = 1.5

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(7)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.5

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(0)
    h3.paragraph_format.space_after = Pt(0)
    h3.paragraph_format.line_spacing = 1.5


def set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_page_number_start(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.append(pg)
    pg.set(qn("w:start"), str(start))


def add_page_field(paragraph, roman: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* ROMAN" if roman else "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, 10.5)


def add_toc_line(doc: Document, title: str, page_hint: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{title}{' …… ' + page_hint if page_hint else ''}")
    set_run_font(run, 12)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Body Text")
    for run in p.runs:
        set_run_font(run, 12)


def add_no_indent_paragraph(doc: Document, text: str, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 12)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.text = ""
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, 16, east_asia="黑体", bold=True)
    elif level == 2:
        set_run_font(run, 14, east_asia="黑体", bold=True)
    else:
        set_run_font(run, 12, east_asia="黑体", bold=True)


def add_formula(doc: Document, formula: str, number: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{formula}    {number}")
    run.font.name = "Cambria Math"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    run.font.size = Pt(12)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 10.5, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(0)
    p.text = ""
    run = p.add_run(str(text))
    set_run_font(run, size, bold=bold)


def shade_cell(cell, fill: str = "F2F2F2") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def add_table_caption(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(caption)
    set_run_font(run, 10.5, east_asia="黑体", bold=True)


def add_figure_caption(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(caption)
    set_run_font(run, 10.5, east_asia="黑体", bold=True)


def add_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str], widths: list[int], max_rows: int | None = None) -> None:
    if df.empty:
        add_no_indent_paragraph(doc, "暂无可用结果。")
        return
    use = df[[c for c in columns if c in df.columns]].copy()
    if max_rows is not None:
        use = use.head(max_rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx])
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for _, row in use.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(use.columns):
            align = WD_ALIGN_PARAGRAPH.LEFT if col in {"province", "target", "model", "script", "description"} else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[idx], fmt_cell(row[col], col), align=align)
    set_table_geometry(table, widths)


def add_figure(doc: Document, rel_path: str, caption: str, width_cm: float = 15.2) -> None:
    path = ROOT / rel_path
    if not path.exists():
        add_no_indent_paragraph(doc, f"图片缺失：{rel_path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_figure_caption(doc, caption)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("期末大作业论文")
    set_run_font(run, 22, east_asia="黑体", bold=True)

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    set_run_font(run, 18, east_asia="黑体", bold=True)

    for _ in range(6):
        doc.add_paragraph()
    rows = [
        ("课程名称", COURSE),
        ("作业题目", TOPIC),
        ("学号", STUDENT_ID),
        ("姓名", STUDENT_NAME),
        ("完成日期", date.today().isoformat()),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2300, 5200])
    for idx, (label, value) in enumerate(rows):
        set_cell_text(table.rows[idx].cells[0], label, bold=True, size=12)
        set_cell_text(table.rows[idx].cells[1], value, size=12, align=WD_ALIGN_PARAGRAPH.LEFT)


def add_abstract(doc: Document, processed: pd.DataFrame) -> None:
    add_heading(doc, "摘要", 1)
    final_confirmed = int(processed["cumulative_confirmed"].dropna().iloc[-1])
    final_deaths = int(processed["cumulative_deaths"].dropna().iloc[-1])
    final_recovered = int(processed["cumulative_recovered"].dropna().iloc[-1])
    add_body_paragraph(
        doc,
        f"本文基于 Johns Hopkins University CSSE COVID-19 全球时间序列数据，对中国疫情传播趋势进行建模与预测研究。"
        f"数据包括累计确诊、累计死亡和累计治愈序列，期末累计确诊 {final_confirmed:,} 例，累计死亡 {final_deaths:,} 例，可靠治愈累计值为 {final_recovered:,} 例。"
        "研究首先完成累计值到每日新增值的转换，并对负新增值、recovered 后期停止更新等数据质量问题进行说明；随后从全国和省份两个层面绘制传播趋势图，补充区域传播趋势分析；"
        "模型部分构建固定参数 SIR、时变参数 SIR 与 GradientBoosting 回归模型，分别完成舱室变量拟合、每日新增预测、未来 7 天和 14 天趋势预测，并使用 RMSE、MAE 和 R² 评价模型表现。"
        "实验结果表明，SIR 类模型具有较强的机制解释能力，时变参数模型能够更好地刻画传播率和移除率变化；GradientBoosting 适合短期每日新增预测，但对数据突变和长期递推预测较敏感。",
    )
    add_no_indent_paragraph(doc, "关键词：疫情传播；SIR 模型；时变参数；GradientBoosting；时间序列预测；区域传播趋势；模型评价")


def add_toc(doc: Document) -> None:
    add_heading(doc, "目录", 1)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    add_toc_field(p)


def build_document() -> None:
    doc = Document()
    configure_styles(doc)
    set_update_fields(doc)

    processed = read_csv("data/processed/china_all_timeseries.csv")
    processed["date"] = pd.to_datetime(processed["date"])
    quality = read_csv("data/processed/china_all_data_quality.csv")
    province = read_csv("outputs/metrics/china_province_trend_summary.csv")
    sir_params = read_csv("outputs/metrics/china_all_sir_parameters.csv")
    sir_metrics = read_csv("outputs/metrics/china_all_sir_metrics.csv")
    tv_params = read_csv("outputs/metrics/china_all_time_varying_sir_parameters.csv")
    tv_metrics = read_csv("outputs/metrics/china_all_time_varying_sir_metrics.csv")
    ml_metrics = read_csv("outputs/metrics/china_all_ml_metrics.csv")
    comp_metrics = read_csv("outputs/metrics/china_all_compartment_ml_metrics.csv")
    future7 = read_csv("outputs/predictions/china_all_ml_future_7d.csv")
    future14 = read_csv("outputs/predictions/china_all_ml_future_14d.csv")

    # Section 1: cover, no footer page number.
    add_cover(doc)

    # Section 2: abstract and TOC, Roman page numbering.
    abstract_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(abstract_section)
    abstract_section.footer.is_linked_to_previous = False
    add_page_field(abstract_section.footer.paragraphs[0], roman=True)
    set_page_number_start(abstract_section, 1)
    add_abstract(doc, processed)
    doc.add_page_break()
    add_toc(doc)

    # Section 3: main body, Arabic page numbering from 1.
    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(body_section)
    body_section.footer.is_linked_to_previous = False
    add_page_field(body_section.footer.paragraphs[0], roman=False)
    set_page_number_start(body_section, 1)

    add_heading(doc, "一、引言", 1)
    add_body_paragraph(doc, "新冠疫情传播数据同时具有传染病动力学特征和时间序列统计特征。累计确诊、死亡和治愈能够反映疫情总体规模，每日新增序列则更能体现短期传播强度、上报波动和防控措施影响。公开疫情数据库和实时监测平台为疫情建模提供了重要数据基础[1-2]。课程大作业不仅要求给出预测曲线，还需要说明模型假设、特征构造、评价指标和结果解释，因此本文采用传统传染病模型与机器学习模型结合的建模思路。")
    add_body_paragraph(doc, "本文主要工作包括：第一，读取 JHU CSSE COVID-19 真实数据并完成预处理；第二，补充全国和省份层面的区域传播趋势分析；第三，实现固定参数 SIR、时变参数 SIR 和 GradientBoosting 回归模型；第四，输出测试集指标、未来 7 天和 14 天预测结果，并将图表与指标整理为论文式文档。")
    add_body_paragraph(doc, "从建模任务角度看，疫情传播预测并不是单纯的曲线外推问题。传染病动力学模型强调传播机制，例如感染者与易感者接触会推动新增感染，治愈和死亡会使感染者转入移除状态；机器学习方法则更关注从历史数据中学习短期变化规律。SIR 及其扩展模型是传染病动力学研究中的经典框架[3-5]，统计学习和集成学习方法则常用于从历史数据中构建预测模型[7-11]。二者的侧重点不同：前者可解释性强，适合说明传播参数含义；后者灵活性较好，适合在已有时间序列基础上进行短期预测。本文将两类方法放在同一个项目中，是为了同时满足“传播规律解释”和“预测效果评价”两方面要求。")
    add_body_paragraph(doc, "本文不追求堆叠复杂模型，而是围绕课程要求构建一个可复现的完整流程。所有图表、指标和预测结果均由程序运行生成，报告文字与输出文件保持一致。对于 recovered 数据停止更新、负新增值、R² 为负等不理想现象，本文不进行人工修饰，而是在结果分析部分如实讨论其原因和影响。")

    add_heading(doc, "二、数据来源与预处理", 1)
    add_heading(doc, "2.1 数据来源", 2)
    start = processed["date"].min().date().isoformat()
    end = processed["date"].max().date().isoformat()
    add_body_paragraph(doc, f"本文使用 Johns Hopkins University Center for Systems Science and Engineering 发布的 COVID-19 global time series 数据[1]。核心文件包括累计确诊、累计死亡和累计治愈三类时间序列。本文分析区域为中国，时间范围为 {start} 至 {end}。")
    add_body_paragraph(doc, "JHU CSSE 数据集是疫情建模研究中常用的公开数据源之一，其优点是覆盖范围广、更新频率高、数据结构相对统一，便于进行跨国家和跨地区的时间序列分析。相关实时可视化研究也证明了该类公开数据在疫情监测中的应用价值[2]。本文选取中国数据作为研究对象，一方面是因为该数据包含省份层级信息，适合补充区域传播趋势分析；另一方面，累计确诊、死亡、治愈三类序列能够支持 SIR 舱室变量的近似构造。")
    add_body_paragraph(doc, "需要注意的是，公开疫情数据通常会受到统计口径、检测能力、上报延迟和历史订正的影响。世界卫生组织和国内公开资料均提示，不同阶段的疫情统计口径和公共卫生措施会影响数据解释[6,12]。尤其是 recovered 序列在部分地区后期不再可靠更新，如果直接将其作为真实治愈数据使用，可能会造成模型输入失真。因此本文在预处理阶段加入数据质量检查，对终端归零或停止更新的 recovered 序列进行标记，并在报告中说明这一限制。")
    add_heading(doc, "2.2 每日新增计算", 2)
    add_formula(doc, "daily_new(t)=max(cumulative(t)-cumulative(t-1),0)", "（1）")
    add_body_paragraph(doc, "式（1）表示累计值到每日新增值的转换。由于真实疫情数据可能出现统计回补或口径修正，累计差分后可能产生负新增值，本文将负值截断为 0，并在数据质量表中记录。")
    add_body_paragraph(doc, "时间序列建模时，训练集和测试集必须按照时间顺序划分，不能随机打乱。随机划分会使模型在训练阶段间接看到未来时期的数据分布，从而造成时间序列预测中的信息泄露。时间序列预测通常强调按时间顺序构造训练和验证流程[10]。本文采用前 80% 样本作为训练集、后 20% 样本作为测试集的基本策略，所有模型评价指标均只在测试集或拟合窗口内计算。")
    add_body_paragraph(doc, "机器学习特征构造遵循“不使用未来信息”的原则。滞后特征只使用预测日前若干天的历史值，滑动均值、滑动标准差和指数滑动均值均基于前一日及以前的数据计算。这样可以保证模型训练过程与真实预测场景一致。")
    add_table_caption(doc, "表 1 数据质量与可靠区间摘要")
    add_table(doc, quality, ["target", "start_date", "end_date", "reliable_end_date", "days", "negative_daily_values_clipped", "missing_days"], ["指标", "开始日期", "结束日期", "可靠结束", "天数", "负新增截断", "缺失天数"], [1100, 1350, 1350, 1350, 900, 1500, 1100])

    add_heading(doc, "三、疫情传播规律与区域趋势分析", 1)
    add_body_paragraph(doc, "全国趋势图用于观察疫情总体规模和阶段性变化，省份趋势图用于比较不同地区之间传播规模与高峰时间差异。为了避免把未分配数据误认为真实省份，JHU 数据中的 Unknown 行仅保留在汇总 CSV 中，不进入 Top 10 省份排序。")
    add_body_paragraph(doc, "从疫情传播规律看，累计曲线通常呈现阶段性增长特征：在传播早期，确诊人数可能快速上升；当防控措施加强、易感接触减少或统计口径变化后，增长速度会发生改变。每日新增曲线比累计曲线更敏感，能够反映短期传播高峰和上报波动，因此本文同时绘制累计趋势和每日新增趋势。")
    add_body_paragraph(doc, "区域传播趋势分析是本题的硬性要求。仅分析全国聚合数据会掩盖不同省份之间的差异，例如部分地区可能出现更早的高峰，部分地区可能在后期增长更明显。本文使用省份面板数据绘制各省累计确诊趋势和每日新增趋势，并通过 Top 10 省份柱状图展示主要贡献地区。这样的区域对比有助于说明疫情传播并非在所有地区同步发生，而是与人口流动、检测能力、局部暴发和统计口径等因素有关。")
    add_figure(doc, "outputs/figures/china_all_cumulative_trends.png", "图 1 全国累计确诊、死亡、治愈趋势")
    add_figure(doc, "outputs/figures/china_all_daily_trends.png", "图 2 全国每日新增确诊、死亡、治愈 7 日均值趋势")
    add_figure(doc, "outputs/figures/china_province_cumulative_confirmed_trends.png", "图 3 各省累计确诊趋势")
    add_figure(doc, "outputs/figures/china_province_daily_confirmed_trends.png", "图 4 各省每日新增确诊 7 日均值趋势")
    add_figure(doc, "outputs/figures/china_province_top10_confirmed_comparison.png", "图 5 Top 10 省份累计确诊对比", width_cm=13.8)
    named = province[province["province"].astype(str).str.upper() != "UNKNOWN"] if not province.empty else province
    add_table_caption(doc, "表 2 Top 10 省份累计确诊与新增高峰")
    add_table(doc, named, ["province", "final_cumulative_confirmed", "total_daily_confirmed", "peak_daily_confirmed"], ["省份", "最终累计确诊", "每日新增合计", "单日新增峰值"], [2200, 2700, 2700, 2500], max_rows=10)

    add_heading(doc, "四、固定参数 SIR 模型", 1)
    add_body_paragraph(doc, "SIR 模型将总人群划分为易感者 S、感染者 I 和移除者 R。本文用累计治愈人数与累计死亡人数之和近似移除者 R，用累计确诊减去移除者近似当前感染者 I。固定参数 SIR 假设传播率 β 和移除率 γ 在拟合窗口内保持不变。")
    add_body_paragraph(doc, "SIR 模型最早用于描述传染病在人群中的传播过程，其基本思想是：易感者与感染者接触后可能被感染，感染者经过隔离、康复或死亡后转入移除状态[3]。模型中的 β 可以理解为单位时间内感染者造成新感染的能力，γ 可以理解为感染者从传播链中移除的速度。虽然真实疫情受到政策干预、检测能力和行为变化影响，但 SIR 模型仍然是理解传播机制的经典基线[4-5]。")
    add_formula(doc, "dS/dt=-βSI/N", "（2）")
    add_formula(doc, "dI/dt=βSI/N-γI", "（3）")
    add_formula(doc, "dR/dt=γI", "（4）")
    add_formula(doc, "R₀=β/γ", "（5）")
    add_body_paragraph(doc, "其中 N 为有效人口规模，β 表示传播率，γ 表示移除率，R₀ 可作为传播强度的近似指标。降低接触率会降低 β，提高隔离和治疗效率会提高 γ，因此二者共同决定传播趋势。")
    add_body_paragraph(doc, "在模型拟合时，本文使用感染者、移除者和累计确诊三类观测量共同约束参数估计，而不是只拟合单一曲线。这样可以减少模型只贴合某一变量而忽略舱室一致性的风险。由于真实数据并不完全符合封闭人群和参数恒定假设，SIR 拟合结果应主要用于解释传播趋势和参数意义，而不应被理解为对真实人口状态的完全还原。")
    add_table_caption(doc, "表 3 固定参数 SIR 拟合参数")
    add_table(doc, sir_params, ["fit_start_date", "fit_end_date", "fit_days", "beta", "gamma", "R0_beta_over_gamma", "effective_population"], ["开始", "结束", "天数", "β", "γ", "R₀", "有效人口"], [1100, 1100, 850, 1100, 1100, 1200, 2000], max_rows=1)
    add_figure(doc, "outputs/figures/china_all_sir_fit.png", "图 6 固定参数 SIR 模型拟合结果")
    add_table_caption(doc, "表 4 固定参数 SIR 模型评价指标")
    add_table(doc, sir_metrics, ["target", "model", "RMSE", "MAE", "R2"], ["目标变量", "模型", "RMSE", "MAE", "R²"], [2600, 1800, 2100, 2100, 1400])

    add_heading(doc, "五、时变参数 SIR 模型", 1)
    add_body_paragraph(doc, "固定参数 SIR 难以描述检测能力、防控政策、医疗资源和公众行为变化带来的阶段差异。时变参数 SIR 将 β 和 γ 扩展为随时间变化的 β(t) 与 γ(t)，用于刻画传播强度和移除效率的动态变化。")
    add_body_paragraph(doc, "疫情传播过程往往不是稳定系统。早期公众防护意识不足、医疗资源紧张、检测能力有限时，传播率和移除率可能与后期明显不同。若仍使用单一 β 和 γ 解释整个窗口，模型可能只能得到一种平均意义上的参数，难以反映传播强度的阶段性变化。因此本文增加时变参数 SIR，尝试用每日参数序列刻画 β(t)、γ(t) 和 Rₜ 的变化。")
    add_formula(doc, "dS/dt=-β(t)SI/N", "（6）")
    add_formula(doc, "dI/dt=β(t)SI/N-γ(t)I", "（7）")
    add_formula(doc, "Rₜ=β(t)/γ(t)", "（8）")
    add_body_paragraph(doc, "与固定参数 SIR 相比，时变参数模型在拟合能力上通常更强，但也更容易受到数据噪声影响。如果 recovered 数据质量较差，或确诊数据存在集中回补，β(t) 与 γ(t) 的估计可能出现波动。因此本文在解释时变参数结果时，更关注其总体变化趋势，而不是过度解读某一天的单点参数。")
    add_table_caption(doc, "表 5 时变参数 SIR 关键参数")
    add_table(doc, tv_params, ["fit_start_date", "fit_end_date", "fit_days", "mean_beta", "mean_gamma", "mean_R0_beta_over_gamma", "final_R0_beta_over_gamma"], ["开始", "结束", "天数", "平均β", "平均γ", "平均Rₜ", "期末Rₜ"], [1050, 1050, 800, 1050, 1050, 1300, 1300], max_rows=1)
    add_figure(doc, "outputs/figures/china_all_time_varying_sir_fit.png", "图 7 时变参数 SIR 模型拟合结果")
    add_table_caption(doc, "表 6 时变参数 SIR 模型评价指标")
    add_table(doc, tv_metrics, ["target", "model", "RMSE", "MAE", "R2"], ["目标变量", "模型", "RMSE", "MAE", "R²"], [2500, 2200, 2000, 2000, 1400])

    add_heading(doc, "六、GradientBoosting 每日新增预测模型", 1)
    add_body_paragraph(doc, "机器学习模型使用 scikit-learn 中的 GradientBoostingRegressor[7]，分别预测每日新增确诊、每日新增死亡和每日新增治愈。特征包括 1 至 14 天及 21、28 天滞后值，3、7、14、28 日滑动均值、滑动中位数、滑动标准差、滑动最大值、指数滑动均值、日期序号、星期和月份等。所有特征均只使用预测日期以前的历史信息，避免未来信息泄露。")
    add_body_paragraph(doc, "GradientBoosting 属于集成学习方法，通过逐步训练多棵弱回归树来拟合前一轮模型的残差[8]。与线性回归相比，它能够刻画非线性关系和特征交互；与深度学习模型相比，它对样本量和计算资源要求较低，适合本课程项目中的中小规模时间序列实验。统计学习教材也将树模型和集成学习视为处理非线性预测任务的重要方法[9,11]。本文保留 GradientBoosting 作为主要机器学习模型，是为了在不引入过重依赖的前提下获得较强的基线预测能力。")
    add_formula(doc, "ŷ(t)=F(x(t);θ), x(t)={lag, rolling, ewm, calendar}", "（9）")
    add_body_paragraph(doc, "对于每日新增序列，单日数据可能受周末效应、集中上报和历史修正影响较大，因此只使用最近一天数据往往不稳定。本文使用多阶滞后、滚动均值、滚动标准差和指数滑动均值等特征，使模型能够同时捕捉短期冲击和较平滑的趋势变化。预测未来多天时，模型采用递推预测：先预测第 1 天，再将预测值加入历史序列用于构造第 2 天特征，以此类推。")
    add_figure(doc, "outputs/figures/china_all_confirmed_ml_test_prediction.png", "图 8 每日新增确诊测试集预测对比")
    add_figure(doc, "outputs/figures/china_all_deaths_ml_test_prediction.png", "图 9 每日新增死亡测试集预测对比")
    add_figure(doc, "outputs/figures/china_all_recovered_ml_test_prediction.png", "图 10 每日新增治愈测试集预测对比")
    add_table_caption(doc, "表 7 GradientBoosting 每日新增预测指标")
    add_table(doc, ml_metrics, ["target", "model", "train_rows", "test_rows", "RMSE", "MAE", "R2"], ["目标", "模型", "训练行", "测试行", "RMSE", "MAE", "R²"], [1200, 2000, 1000, 1000, 1600, 1600, 1200])
    add_figure(doc, "outputs/figures/china_all_ml_metrics_comparison.png", "图 11 GradientBoosting 每日新增预测指标对比", width_cm=14.5)

    add_heading(doc, "七、同窗口舱室变量对比实验", 1)
    add_body_paragraph(doc, "为了让传统模型和机器学习模型在同一类变量上可比，本文将 GradientBoosting 应用于 SIR 相同观测窗口内的 infected、removed 和 confirmed 三个目标。该实验仍按时间顺序划分训练集和测试集。实验中 R² 为负的结果被如实保留，说明机器学习模型在该短窗口后段上表现差于均值基线。")
    add_body_paragraph(doc, "该实验的意义不在于证明 GradientBoosting 一定优于 SIR，而是为论文提供更公平的横向比较。传统模型主要处理舱室变量，而每日新增预测属于另一类任务。如果只比较 SIR 的舱室拟合与 GradientBoosting 的每日新增预测，二者目标变量不同，结论容易混淆。因此本文增加同窗口舱室实验，使机器学习模型也面对 infected、removed 和 confirmed 三个舱室变量。")
    add_body_paragraph(doc, "从结果看，GradientBoosting 在该短窗口舱室变量测试集上的 R² 为负，说明它在这一任务中没有超过均值基线。这一结果并不意味着机器学习方法无效，而是说明当样本窗口较短、趋势后段变化明显、且缺少机制约束时，树模型可能难以稳定外推。相比之下，SIR 模型虽然假设较强，但舱室结构本身提供了额外约束，因此在舱室变量解释上更有优势。")
    add_figure(doc, "outputs/figures/china_all_compartment_ml_test_prediction.png", "图 12 GradientBoosting 舱室变量同窗口测试预测")
    add_table_caption(doc, "表 8 GradientBoosting 舱室变量同窗口实验指标")
    add_table(doc, comp_metrics, ["target", "model", "train_rows", "test_rows", "RMSE", "MAE", "R2"], ["目标", "模型", "训练行", "测试行", "RMSE", "MAE", "R²"], [1200, 2000, 1000, 1000, 1600, 1600, 1200])

    add_heading(doc, "八、未来 7 天和 14 天预测", 1)
    add_body_paragraph(doc, "未来预测采用递推方式完成：每预测一天，就将该预测值加入历史序列，再构造下一天的滞后和滑动统计特征。递推预测是时间序列多步预测中的常见策略，但预测期越长，误差累积风险越明显[10]。")
    add_body_paragraph(doc, "未来 7 天预测通常可视为短期预测，主要反映模型根据最近历史变化推断出的延续趋势；未来 14 天预测则更容易受到误差累积影响。如果第 1 天或第 2 天预测存在偏差，后续滚动特征会继续使用这些预测值，导致偏差逐步传递。因此本文同时输出 7 天和 14 天预测，并在结论中强调其适用范围。")
    add_formula(doc, "RMSE=sqrt((1/n)Σ(yi-ŷi)²)", "（10）")
    add_formula(doc, "MAE=(1/n)Σ|yi-ŷi|", "（11）")
    add_formula(doc, "R²=1-Σ(yi-ŷi)²/Σ(yi-ȳ)²", "（12）")
    add_figure(doc, "outputs/figures/china_all_confirmed_ml_future_7d.png", "图 13 每日新增确诊未来 7 天预测")
    add_figure(doc, "outputs/figures/china_all_confirmed_ml_future_14d.png", "图 14 每日新增确诊未来 14 天预测")
    add_table_caption(doc, "表 9 未来 7 天预测结果")
    add_table(doc, future7, ["target", "model", "date", "forecast_day", "prediction"], ["目标", "模型", "日期", "预测天数", "预测值"], [1300, 2200, 1800, 1300, 2100], max_rows=21)
    add_table_caption(doc, "表 10 未来 14 天预测结果")
    add_table(doc, future14, ["target", "model", "date", "forecast_day", "prediction"], ["目标", "模型", "日期", "预测天数", "预测值"], [1300, 2200, 1800, 1300, 2100], max_rows=42)

    add_heading(doc, "九、模型对比分析", 1)
    add_body_paragraph(doc, "固定参数 SIR 的优点是结构清楚、参数含义明确，能够直观说明接触率降低和治愈率提高对传播控制的作用；缺点是难以覆盖政策和行为变化。时变参数 SIR 更能贴合阶段性变化，但对数据质量和参数估计稳定性要求更高。GradientBoosting 能利用滞后特征进行短期预测，但在突发统计口径变化和长期递推预测中并不稳健。")
    add_body_paragraph(doc, "从评价指标看，每日新增确诊预测具有一定效果，而死亡和治愈序列受数据稀疏性和 recovered 停止更新影响较大。同窗口舱室变量实验中 GradientBoosting 的 R² 为负，说明短窗口舱室变量更适合使用具有机制约束的 SIR 类模型解释。")
    add_body_paragraph(doc, "RMSE、MAE 和 R² 从不同角度反映模型误差。RMSE 对较大误差更敏感，适合观察模型是否在某些高峰期出现较大偏差；MAE 表示平均绝对误差，更直观地反映日均偏离程度；R² 衡量模型相对于均值基线的解释能力。这些指标也是回归和预测任务中常用的评价方式[7,10-11]。当 R² 为负时，说明模型预测效果差于简单使用均值预测。本文保留负 R²，是因为它能够真实反映模型在部分任务上的不足。")
    add_body_paragraph(doc, "综合来看，传统模型和机器学习模型并不是替代关系，而是互补关系。SIR 类模型适合回答“为什么传播会上升或下降”，GradientBoosting 更适合回答“在近期历史模式延续的情况下，短期新增可能是多少”。在真实疫情分析中，若要进一步提升预测能力，还应加入政策强度、检测量、人口流动、疫苗接种等外生变量。")

    add_heading(doc, "十、结论与不足", 1)
    add_body_paragraph(doc, "本文完成了一个完整的疫情传播建模工程：从真实疫情数据读取、累计值到每日新增值转换、区域传播趋势分析，到固定参数 SIR、时变参数 SIR、GradientBoosting 每日新增预测和同窗口舱室变量实验，最终形成图表、指标、预测 CSV 和论文式 Word 文档。")
    add_body_paragraph(doc, "本文仍存在以下不足：第一，JHU 数据存在回补、修正和 recovered 后期停止更新问题；第二，模型未显式加入政策干预、人口流动、检测能力等外生变量；第三，未来预测采用递推方式，长期预测误差会累积；第四，省份之间人口基数和统计口径不同，累计确诊 Top 10 只能反映原始规模，不能直接等同于风险率。")
    add_body_paragraph(doc, "后续工作可以从三个方向改进：一是引入人口规模进行标准化，使用每百万人确诊数或死亡数比较区域风险；二是加入外生变量，例如检测量、节假日、人口迁移和防控政策指标；三是尝试更系统的模型比较，例如 ARIMA、Prophet、XGBoost、LSTM 或 GRU，但仍应保留传统模型作为可解释基线。")

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Johns Hopkins University Center for Systems Science and Engineering. COVID-19 Data Repository, 2020-2023.",
        "[2] Dong E, Du H, Gardner L. An interactive web-based dashboard to track COVID-19 in real time. The Lancet Infectious Diseases, 2020, 20(5): 533-534.",
        "[3] Kermack W O, McKendrick A G. A contribution to the mathematical theory of epidemics. Proceedings of the Royal Society A, 1927, 115(772): 700-721.",
        "[4] Hethcote H W. The mathematics of infectious diseases. SIAM Review, 2000, 42(4): 599-653.",
        "[5] Anderson R M, May R M. Infectious Diseases of Humans: Dynamics and Control. Oxford University Press, 1991.",
        "[6] World Health Organization. Coronavirus disease situation reports and public health guidance, 2020-2023.",
        "[7] Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 2011, 12: 2825-2830.",
        "[8] Friedman J H. Greedy function approximation: A gradient boosting machine. Annals of Statistics, 2001, 29(5): 1189-1232.",
        "[9] Hastie T, Tibshirani R, Friedman J. The Elements of Statistical Learning. Springer, 2009.",
        "[10] Hyndman R J, Athanasopoulos G. Forecasting: Principles and Practice. OTexts, 2021.",
        "[11] James G, Witten D, Hastie T, Tibshirani R. An Introduction to Statistical Learning. Springer, 2021.",
        "[12] 中华人民共和国国家卫生健康委员会. 新型冠状病毒感染疫情防控相关公开资料, 2020-2023.",
    ]
    for ref in refs:
        add_no_indent_paragraph(doc, ref)

    add_heading(doc, "附录：核心代码说明", 1)
    scripts = pd.DataFrame(
        [
            ("src/prepare_data.py", "读取 JHU 数据，生成全国聚合序列和省份面板。"),
            ("src/run_regional_trends.py", "生成全国趋势图、省份趋势图和 Top 10 省份对比图。"),
            ("src/run_sir.py", "拟合固定参数 SIR，输出参数、指标和拟合图。"),
            ("src/run_time_varying_sir.py", "拟合时变参数 SIR，输出每日参数和拟合图。"),
            ("src/run_ml.py", "训练 GradientBoosting 每日新增预测模型并输出 7/14 天预测。"),
            ("src/run_compartment_ml.py", "在 SIR 同窗口上对 infected、removed、confirmed 做机器学习对比。"),
        ],
        columns=["script", "description"],
    )
    add_table_caption(doc, "表 11 核心脚本与功能对应关系")
    add_table(doc, scripts, ["script", "description"], ["脚本", "功能"], [3300, 6300])

    DEST.parent.mkdir(exist_ok=True)
    doc.save(DEST)


if __name__ == "__main__":
    build_document()
