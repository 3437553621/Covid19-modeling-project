from __future__ import annotations

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "1030424322-解宝赛" / "期末大作业论文.docx"
REPORT_COPY = ROOT / "1030424322-解宝赛" / "reports" / "report.md"

STUDENT_ID = "1030424322"
STUDENT_NAME = "解宝赛"
COURSE = "计算机建模技术"
TITLE = "基于传染病动力学模型与机器学习方法的疫情传播预测研究"
SUBTITLE = "题目一：疫情传播模型与预测模型方法"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(89, 89, 89)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"
BORDER = "B7C9DD"


def fmt(value, digits: int = 3) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, (float, int)):
        value = float(value)
        if abs(value) >= 1000:
            return f"{value:,.{digits}f}"
        return f"{value:.{digits}f}"
    return str(value)


def read_csv(path: str | Path) -> pd.DataFrame:
    path = ROOT / path
    return pd.read_csv(path) if path.exists() else pd.DataFrame()


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    p.text = ""
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    run.bold = bold


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in [("top", 80), ("bottom", 80), ("start", 120), ("end", 120)]:
                elem = tc_mar.find(qn(f"w:{side}"))
                if elem is None:
                    elem = OxmlElement(f"w:{side}")
                    tc_mar.append(elem)
                elem.set(qn("w:w"), str(value))
                elem.set(qn("w:type"), "dxa")


def add_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str], widths: list[int], max_rows: int | None = None) -> None:
    if df.empty:
        p = doc.add_paragraph("暂无可用结果。")
        p.style = "Body Text"
        return
    use = df.copy()
    use = use[[c for c in columns if c in use.columns]]
    if max_rows is not None:
        use = use.head(max_rows)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], HEADER_FILL)
        set_cell_text(table.rows[0].cells[idx], header, bold=True)

    for _, row in use.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(use.columns):
            align = WD_ALIGN_PARAGRAPH.LEFT if col in {"target", "model", "province", "evaluation_scope"} else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[idx], fmt(row[col]), align=align)
    set_table_geometry(table, widths)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9.5)
    run.font.color.rgb = GRAY
    run.bold = True


def add_figure(doc: Document, rel_path: str, caption: str, width: float = 6.1) -> None:
    path = ROOT / rel_path
    if not path.exists():
        p = doc.add_paragraph(f"图片缺失：{rel_path}")
        p.style = "Body Text"
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_formula(doc: Document, text: str, number: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.text = ""
    run = p.add_run(f"{text}    {number}")
    run.font.name = "Cambria Math"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    run.font.size = Pt(11)


def add_paragraph(doc: Document, text: str, style: str = "Body Text") -> None:
    p = doc.add_paragraph(text)
    p.style = style


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    paragraph.add_run(" 页")


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    body = styles["Body Text"]
    body.font.name = "Calibri"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    body.font.size = Pt(11)
    body.paragraph_format.first_line_indent = Pt(22)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.space_after = Pt(8)
    body.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("期末大作业论文")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_BLUE
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(TITLE)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SUBTITLE)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(13)
    run.font.color.rgb = GRAY

    for _ in range(5):
        doc.add_paragraph()

    rows = [
        ("课程名称", COURSE),
        ("学号", STUDENT_ID),
        ("姓名", STUDENT_NAME),
        ("数据来源", "Johns Hopkins University CSSE COVID-19 Data"),
        ("完成日期", date.today().isoformat()),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2200, 5000])
    for idx, (label, value) in enumerate(rows):
        shade_cell(table.rows[idx].cells[0], HEADER_FILL)
        set_cell_text(table.rows[idx].cells[0], label, bold=True)
        set_cell_text(table.rows[idx].cells[1], value, align=WD_ALIGN_PARAGRAPH.LEFT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    run = p.add_run("提交说明：本论文对应程序、实验结果和图表均保存在同名项目文件夹中。")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY
    doc.add_page_break()


def add_static_toc(doc: Document) -> None:
    doc.add_heading("目录", level=1)
    items = [
        "摘要",
        "一、引言",
        "二、数据来源与预处理",
        "三、疫情传播规律与区域趋势分析",
        "四、固定参数 SIR 模型",
        "五、时变参数 SIR 模型",
        "六、GradientBoosting 每日新增预测模型",
        "七、同窗口舱室变量公平对比实验",
        "八、未来 7 天和 14 天趋势预测",
        "九、模型评价与综合讨论",
        "十、结论与不足",
        "参考文献",
        "附录：核心代码与文件说明",
    ]
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(3)
    doc.add_page_break()


def build_document() -> None:
    dest_dir = ROOT / "1030424322-解宝赛"
    dest_dir.mkdir(exist_ok=True)
    (dest_dir / "reports").mkdir(exist_ok=True)

    doc = Document()
    configure_styles(doc)
    doc.sections[0].header.paragraphs[0].text = TITLE
    doc.sections[0].header.paragraphs[0].style = doc.styles["Normal"]
    doc.sections[0].footer.paragraphs[0].text = ""
    add_page_number(doc.sections[0].footer.paragraphs[0])

    add_cover(doc)
    add_static_toc(doc)

    processed = read_csv("data/processed/china_all_timeseries.csv")
    quality = read_csv("data/processed/china_all_data_quality.csv")
    province = read_csv("outputs/metrics/china_province_trend_summary.csv")
    sir_params = read_csv("outputs/metrics/china_all_sir_parameters.csv")
    sir_metrics = read_csv("outputs/metrics/china_all_sir_metrics.csv")
    tv_params = read_csv("outputs/metrics/china_all_time_varying_sir_parameters.csv")
    tv_metrics = read_csv("outputs/metrics/china_all_time_varying_sir_metrics.csv")
    ml_metrics = read_csv("outputs/metrics/china_all_ml_metrics.csv")
    comp_metrics = read_csv("outputs/metrics/china_all_compartment_ml_metrics.csv")
    future7 = read_csv("outputs/predictions/china_all_ml_future_7d.csv")
    future14 = read_csv("outputs/predictions/china_all_ml_future_14d.csv")

    processed["date"] = pd.to_datetime(processed["date"])
    start = processed["date"].min().date().isoformat()
    end = processed["date"].max().date().isoformat()
    final_confirmed = int(processed["cumulative_confirmed"].dropna().iloc[-1])
    final_deaths = int(processed["cumulative_deaths"].dropna().iloc[-1])
    final_recovered = int(processed["cumulative_recovered"].dropna().iloc[-1])

    doc.add_heading("摘要", level=1)
    add_paragraph(
        doc,
        f"本文围绕《计算机建模技术》期末大作业题目一，基于 Johns Hopkins University CSSE COVID-19 全球时间序列数据，对中国疫情传播趋势进行建模与预测研究。"
        f"数据时间范围为 {start} 至 {end}，期末累计确诊 {final_confirmed:,} 例、累计死亡 {final_deaths:,} 例、可靠治愈累计值截至数据停止更新前为 {final_recovered:,} 例。"
        "研究首先完成累计数据到每日新增数据的转换，并对负新增值和 recovered 后期停止更新问题进行说明；随后从全国与省份两个层面绘制趋势图，分析区域传播差异；"
        "模型部分实现固定参数 SIR、时变参数 SIR 以及 GradientBoosting 回归模型，并分别输出测试集预测、未来 7 天和 14 天预测、RMSE、MAE、R2 等指标。"
        "实验结果表明，SIR 类模型具有较强的机制解释能力，时变参数模型能够更灵活地描述传播率和移除率变化；GradientBoosting 适合短期时间序列预测，但在突发统计口径变化和长期递推预测中仍存在误差累积风险。",
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.add_run("关键词：").bold = True
    p.add_run("疫情传播；SIR 模型；时变参数；GradientBoosting；时间序列预测；区域传播趋势；模型评价")

    doc.add_heading("一、引言", level=1)
    add_paragraph(
        doc,
        "新冠疫情传播数据兼具传染病动力学特征和时间序列统计特征。累计确诊、死亡、治愈等指标能够反映疫情总体规模，而每日新增序列更能体现短期传播强度和上报波动。"
        "对于课程建模任务而言，仅给出预测曲线并不足够，还需要说明模型假设、数据处理方法、评价指标以及区域差异。本文因此采用“机制模型 + 机器学习模型”的组合思路："
        "一方面通过 SIR 模型刻画易感者、感染者和移除者之间的转化关系，另一方面通过 GradientBoosting 利用滞后特征和滑动统计量完成短期预测。",
    )
    add_paragraph(
        doc,
        "本文的工作重点包括四个方面。第一，基于 JHU 数据完成可复现的数据处理流程；第二，补充省份层面区域传播趋势分析；第三，在相同舱室观测窗口下比较传统模型和机器学习模型；"
        "第四，将指标、预测结果和图表整理成论文式报告，保证结论与程序输出一致。",
    )

    doc.add_heading("二、数据来源与预处理", level=1)
    add_paragraph(
        doc,
        "本文使用 Johns Hopkins University Center for Systems Science and Engineering 发布的 COVID-19 global time series 数据。"
        "核心文件包括累计确诊、累计死亡和累计治愈三类宽表。每个文件前四列为地区元数据，后续列为日期。程序首先按国家筛选中国数据，若未指定省份则对同一国家的所有省份行进行求和。"
    )
    add_formula(doc, "daily_new(t) = max(cumulative(t) - cumulative(t - 1), 0)", "（1）")
    add_paragraph(
        doc,
        "式（1）表示累计值到每日新增值的转换。由于真实疫情数据可能发生历史订正和统计回补，累计差分后会出现负新增。本文将负新增截断为 0，并在报告中明确说明其含义。"
        "训练集和测试集按时间顺序划分，避免随机打乱时间序列造成未来信息泄露。",
    )
    add_caption(doc, "表 1 数据质量与可靠区间摘要")
    add_table(
        doc,
        quality,
        ["target", "start_date", "end_date", "reliable_end_date", "days", "negative_daily_values_clipped", "missing_days", "possible_reset_or_stop"],
        ["指标", "开始日期", "结束日期", "可靠结束", "天数", "负新增截断", "缺失天数", "可能停止更新"],
        [1100, 1200, 1200, 1300, 800, 1400, 1100, 1250],
    )

    doc.add_heading("三、疫情传播规律与区域趋势分析", level=1)
    add_paragraph(
        doc,
        "全国趋势图用于观察疫情总体规模和阶段性变化。累计曲线反映总量变化，每日新增曲线使用 7 日均值弱化单日上报波动。"
        "省份层面图则用于满足题目中的“区域传播趋势”要求，展示不同地区之间传播规模和高峰时间的差异。",
    )
    add_figure(doc, "outputs/figures/china_all_cumulative_trends.png", "图 1 全国累计确诊、死亡、治愈趋势")
    add_figure(doc, "outputs/figures/china_all_daily_trends.png", "图 2 全国每日新增确诊、死亡、治愈 7 日均值趋势")
    add_figure(doc, "outputs/figures/china_province_cumulative_confirmed_trends.png", "图 3 各省累计确诊趋势")
    add_figure(doc, "outputs/figures/china_province_daily_confirmed_trends.png", "图 4 各省每日新增确诊 7 日均值趋势")
    add_figure(doc, "outputs/figures/china_province_top10_confirmed_comparison.png", "图 5 Top 10 省份累计确诊对比")
    add_paragraph(
        doc,
        "JHU 数据中的 Unknown 表示未分配地区，不属于严格意义上的省份。为保证区域对比的可解释性，本文在省份汇总 CSV 中保留 Unknown 以保持数据透明，但在 Top 10 省份图和下表中将其排除。",
    )
    province_named = province[province["province"].astype(str).str.upper() != "UNKNOWN"] if not province.empty else province
    add_caption(doc, "表 2 Top 10 省份累计确诊与新增高峰")
    add_table(
        doc,
        province_named,
        ["province", "final_cumulative_confirmed", "total_daily_confirmed", "peak_daily_confirmed"],
        ["省份", "最终累计确诊", "每日新增合计", "单日新增峰值"],
        [1900, 2500, 2500, 2300],
        max_rows=10,
    )

    doc.add_heading("四、固定参数 SIR 模型", level=1)
    add_paragraph(
        doc,
        "SIR 模型将总人群划分为易感者 S、感染者 I 和移除者 R。本文用累计治愈人数与累计死亡人数之和近似移除者 R，用累计确诊减去移除者近似当前感染者 I。"
        "固定参数 SIR 假设传播率 beta 和移除率 gamma 在拟合窗口内保持不变，适合解释基本传播机制。",
    )
    add_formula(doc, "dS/dt = -βSI/N", "（2）")
    add_formula(doc, "dI/dt = βSI/N - γI", "（3）")
    add_formula(doc, "dR/dt = γI", "（4）")
    add_formula(doc, "R0 = β / γ", "（5）")
    add_paragraph(
        doc,
        "其中 N 为有效人口规模，β 表示单位时间传播率，γ 表示感染者转入移除状态的速率。β 降低通常对应接触率下降或防控措施加强，γ 提高通常对应隔离、治疗和康复效率提升。"
        "R0 可作为传播强度的近似指标，数值越高表示传播扩张能力越强。",
    )
    add_caption(doc, "表 3 固定参数 SIR 拟合参数")
    add_table(
        doc,
        sir_params,
        ["fit_start_date", "fit_end_date", "fit_days", "beta", "gamma", "R0_beta_over_gamma", "effective_population"],
        ["开始", "结束", "天数", "β", "γ", "R0", "有效人口"],
        [1100, 1100, 800, 1000, 1000, 1200, 2100],
        max_rows=1,
    )
    add_figure(doc, "outputs/figures/china_all_sir_fit.png", "图 6 固定参数 SIR 模型拟合结果")
    add_caption(doc, "表 4 固定参数 SIR 模型评价指标")
    add_table(
        doc,
        sir_metrics,
        ["target", "model", "RMSE", "MAE", "R2"],
        ["目标变量", "模型", "RMSE", "MAE", "R2"],
        [2200, 1700, 1800, 1800, 1300],
    )

    doc.add_heading("五、时变参数 SIR 模型", level=1)
    add_paragraph(
        doc,
        "固定参数模型难以描述检测能力、隔离政策、医疗资源和公众行为变化带来的阶段差异。时变参数 SIR 将 β 和 γ 扩展为随时间变化的 β(t) 与 γ(t)，用于刻画传播强度和移除效率的动态变化。",
    )
    add_formula(doc, "dS/dt = -β(t)SI/N", "（6）")
    add_formula(doc, "dI/dt = β(t)SI/N - γ(t)I", "（7）")
    add_formula(doc, "Rt = β(t) / γ(t)", "（8）")
    add_caption(doc, "表 5 时变参数 SIR 关键参数")
    add_table(
        doc,
        tv_params,
        ["fit_start_date", "fit_end_date", "fit_days", "mean_beta", "mean_gamma", "mean_R0_beta_over_gamma", "final_R0_beta_over_gamma"],
        ["开始", "结束", "天数", "平均β", "平均γ", "平均Rt", "期末Rt"],
        [1100, 1100, 800, 1100, 1100, 1400, 1400],
        max_rows=1,
    )
    add_figure(doc, "outputs/figures/china_all_time_varying_sir_fit.png", "图 7 时变参数 SIR 模型拟合结果")
    add_caption(doc, "表 6 时变参数 SIR 模型评价指标")
    add_table(
        doc,
        tv_metrics,
        ["target", "model", "RMSE", "MAE", "R2"],
        ["目标变量", "模型", "RMSE", "MAE", "R2"],
        [2200, 2100, 1700, 1700, 1300],
    )

    doc.add_heading("六、GradientBoosting 每日新增预测模型", level=1)
    add_paragraph(
        doc,
        "机器学习模型部分使用 GradientBoostingRegressor。该模型通过多棵弱回归树逐步拟合残差，能够处理非线性关系和特征交互。"
        "本文针对每日新增确诊、死亡、治愈分别建模。特征来自历史序列本身，包括 1 至 14 天以及 21、28 天滞后值，3、7、14、28 日滑动均值、滑动中位数、滑动标准差、滑动最大值、指数滑动均值、日期序号、星期和月份等。",
    )
    add_formula(doc, "ŷ(t) = F(x(t); θ),  x(t) = {lag, rolling, ewm, calendar features}", "（9）")
    add_paragraph(
        doc,
        "所有监督学习特征均只使用预测日期之前的历史信息构造，不使用未来真实值。测试集按时间顺序取后 20%，最终预测值使用 max(ŷ, 0) 截断为非负数。",
    )
    add_figure(doc, "outputs/figures/china_all_confirmed_ml_test_prediction.png", "图 8 每日新增确诊测试集预测对比")
    add_figure(doc, "outputs/figures/china_all_deaths_ml_test_prediction.png", "图 9 每日新增死亡测试集预测对比")
    add_figure(doc, "outputs/figures/china_all_recovered_ml_test_prediction.png", "图 10 每日新增治愈测试集预测对比")
    add_caption(doc, "表 7 GradientBoosting 每日新增预测指标")
    add_table(
        doc,
        ml_metrics,
        ["target", "model", "train_rows", "test_rows", "RMSE", "MAE", "R2"],
        ["目标", "模型", "训练行", "测试行", "RMSE", "MAE", "R2"],
        [1200, 1900, 1000, 1000, 1500, 1500, 1260],
    )
    add_figure(doc, "outputs/figures/china_all_ml_metrics_comparison.png", "图 11 GradientBoosting 每日新增预测指标对比")

    doc.add_heading("七、同窗口舱室变量公平对比实验", level=1)
    add_paragraph(
        doc,
        "为了避免传统模型和机器学习模型只在不同目标上比较，本文增加 GradientBoosting 在 SIR 相同窗口上的舱室变量预测实验。"
        "该实验以 SIR 观测窗口内的 infected、removed、confirmed 为目标，仍采用滞后和滑动统计特征，并按时间顺序划分训练集与测试集。这样论文中可以同时说明：SIR 和机器学习模型均在同一类舱室数据上进行过对比，而 GradientBoosting 还额外承担每日新增预测任务。",
    )
    add_figure(doc, "outputs/figures/china_all_compartment_ml_test_prediction.png", "图 12 GradientBoosting 舱室变量同窗口测试预测")
    add_caption(doc, "表 8 GradientBoosting 舱室变量同窗口实验指标")
    add_table(
        doc,
        comp_metrics,
        ["target", "model", "train_rows", "test_rows", "RMSE", "MAE", "R2"],
        ["目标", "模型", "训练行", "测试行", "RMSE", "MAE", "R2"],
        [1200, 1900, 1000, 1000, 1500, 1500, 1260],
    )
    add_paragraph(
        doc,
        "从表 8 可以看到，该同窗口实验的 R2 为负，说明 GradientBoosting 在 SIR 短窗口测试段上的表现差于均值基线。该结果没有被删除或修饰，而是作为模型适用性讨论的一部分："
        "机制模型在舱室变量拟合上更具结构优势，而树模型在样本窗口较短、后段趋势突变时可能难以稳定外推。",
    )

    doc.add_heading("八、未来 7 天和 14 天趋势预测", level=1)
    add_paragraph(
        doc,
        "未来预测采用递推方式完成：每预测一天，就将该预测值加入历史序列，再构造下一天的滞后和滑动特征。该方法适合短期趋势展示，但预测期越长，误差累积风险越明显。",
    )
    add_figure(doc, "outputs/figures/china_all_confirmed_ml_future_7d.png", "图 13 每日新增确诊未来 7 天预测")
    add_figure(doc, "outputs/figures/china_all_confirmed_ml_future_14d.png", "图 14 每日新增确诊未来 14 天预测")
    add_caption(doc, "表 9 未来 7 天预测结果")
    add_table(
        doc,
        future7,
        ["target", "model", "date", "forecast_day", "prediction"],
        ["目标", "模型", "日期", "预测天数", "预测值"],
        [1300, 2100, 1700, 1300, 2200],
        max_rows=21,
    )
    add_caption(doc, "表 10 未来 14 天预测结果")
    add_table(
        doc,
        future14,
        ["target", "model", "date", "forecast_day", "prediction"],
        ["目标", "模型", "日期", "预测天数", "预测值"],
        [1300, 2100, 1700, 1300, 2200],
        max_rows=42,
    )

    doc.add_heading("九、模型评价与综合讨论", level=1)
    add_paragraph(
        doc,
        "本文使用 RMSE、MAE 与 R2 作为主要评价指标。RMSE 对大误差更敏感，MAE 表示平均绝对偏差，R2 用于衡量模型相对均值基线的解释能力。"
        "当 R2 为负时，说明模型在测试集上不如均值预测，本文保留这些结果以反映真实模型局限。",
    )
    add_formula(doc, "RMSE = sqrt((1/n)Σ(yi - ŷi)^2)", "（10）")
    add_formula(doc, "MAE = (1/n)Σ|yi - ŷi|", "（11）")
    add_formula(doc, "R² = 1 - Σ(yi - ŷi)^2 / Σ(yi - ȳ)^2", "（12）")
    add_paragraph(
        doc,
        "固定参数 SIR 的优点是结构清晰、参数可解释，能够直观说明接触率降低和治愈率提高对传播控制的作用；缺点是难以覆盖政策变化和统计口径变化。"
        "时变参数 SIR 更能贴合阶段性变化，但也更依赖数据质量和参数估计稳定性。GradientBoosting 对每日新增确诊的短期预测具有一定效果，但对死亡和治愈序列的预测受数据质量、稀疏性和 recovered 停止更新影响更大。",
    )
    add_bullets(
        doc,
        [
            "机制解释：SIR/时变 SIR 能解释传播率、移除率和传播强度变化。",
            "短期预测：GradientBoosting 能利用滞后特征捕捉短期变化，但不擅长长期递推。",
            "区域分析：省份趋势图揭示不同地区规模差异，Top 10 对比可辅助发现重点地区。",
            "评价原则：所有指标均由程序生成，负 R2 如实保留，不进行人工修饰。",
        ],
    )

    doc.add_heading("十、结论与不足", level=1)
    add_paragraph(
        doc,
        "本文完成了一个较完整的疫情传播建模工程：从 JHU 真实数据读取、累计值到每日新增值转换、区域传播趋势图绘制，到固定参数 SIR、时变参数 SIR、GradientBoosting 每日新增预测和同窗口舱室变量实验，最终形成图表、指标、预测 CSV 和论文式 Word 文档。"
        "从建模结果看，传统传染病模型在解释传播机制方面更有优势，机器学习模型更适合基于历史特征进行短期预测；二者结合能够同时满足课程对模型方法、评价指标和可视化分析的要求。",
    )
    add_paragraph(
        doc,
        "本文仍存在以下不足：第一，JHU 数据存在回补、修正和 recovered 后期停止更新问题；第二，模型未显式加入政策干预、人口流动、检测能力等外生变量；第三，未来预测采用递推方式，长期预测误差会累积；"
        "第四，省份之间人口基数和统计口径不同，累计确诊 Top 10 只能反映原始规模，不能直接等同于风险率；第五，本文保留模型效果较差的结果用于真实讨论，但仍可在后续工作中引入更多外生变量或更稳健的时间序列模型进行改进。",
    )

    doc.add_heading("参考文献", level=1)
    refs = [
        "Johns Hopkins University Center for Systems Science and Engineering. COVID-19 Data Repository.",
        "Kermack W O, McKendrick A G. A contribution to the mathematical theory of epidemics.",
        "scikit-learn developers. Gradient Boosting regression documentation.",
        "World Health Organization. Coronavirus disease situation reports and public health guidance.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style="List Number")
        p.paragraph_format.space_after = Pt(4)

    doc.add_heading("附录：核心代码与文件说明", level=1)
    add_caption(doc, "表 11 核心脚本与功能对应关系")
    scripts = pd.DataFrame(
        [
            ("src/prepare_data.py", "读取 JHU 数据，生成全国聚合序列和省份面板。"),
            ("src/run_regional_trends.py", "生成全国趋势图、省份趋势图和 Top 10 省份对比图。"),
            ("src/run_sir.py", "拟合固定参数 SIR，输出参数、指标和拟合图。"),
            ("src/run_time_varying_sir.py", "拟合时变参数 SIR，输出每日参数和拟合图。"),
            ("src/run_ml.py", "训练 GradientBoosting 每日新增预测模型并输出 7/14 天预测。"),
            ("src/run_compartment_ml.py", "在 SIR 同窗口上对 infected、removed、confirmed 做机器学习对比。"),
            ("src/generate_report.py", "生成 Markdown 论文初稿。"),
            ("scripts/build_paper_docx.py", "生成包含公式、表格和图片的 Word 论文。"),
        ],
        columns=["script", "description"],
    )
    add_table(doc, scripts, ["script", "description"], ["脚本", "功能"], [2500, 6860])

    OUTPUT.parent.mkdir(exist_ok=True)
    doc.save(OUTPUT)

    report_source = ROOT / "reports" / "report.md"
    if report_source.exists():
        REPORT_COPY.write_text(report_source.read_text(encoding="utf-8"), encoding="utf-8")


if __name__ == "__main__":
    build_document()
